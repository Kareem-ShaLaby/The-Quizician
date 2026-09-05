import re
import string
import random
import json
import os
import asyncio
import html
import tempfile
from io import BytesIO


from telegram import Update, ReactionTypeEmoji, InlineKeyboardButton, InlineKeyboardMarkup, InputMediaPhoto, InputFile
from telegram.error import Forbidden
from telegram.ext import (
    ApplicationBuilder,
    MessageHandler,
    CommandHandler,
    CallbackQueryHandler,
    PollHandler,
    PollAnswerHandler,
    filters,
    ContextTypes,
)
from telegram.constants import ParseMode

from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Image as RLImage
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    # python-docx (and its lxml dependency) not installed — DOCX export is
    # simply disabled until it's installed; everything else works fine.
    DOCX_AVAILABLE = False

# ═══════════════════════════════════════════════════════════════
# FONT SETUP
# ═══════════════════════════════════════════════════════════════
_POPPINS_REG  = "/usr/share/fonts/truetype/google-fonts/Poppins-Regular.ttf"
_POPPINS_BOLD = "/usr/share/fonts/truetype/google-fonts/Poppins-Bold.ttf"

FONT_NAME      = "Helvetica"
FONT_NAME_BOLD = "Helvetica-Bold"

if os.path.exists(_POPPINS_REG) and os.path.exists(_POPPINS_BOLD):
    try:
        pdfmetrics.registerFont(TTFont("Poppins",      _POPPINS_REG))
        pdfmetrics.registerFont(TTFont("Poppins-Bold", _POPPINS_BOLD))
        FONT_NAME      = "Poppins"
        FONT_NAME_BOLD = "Poppins-Bold"
        print("Poppins font loaded")
    except Exception as e:
        print(f"Poppins load error: {e} — using Helvetica")
else:
    print("Poppins not found — using Helvetica")

BOT_TOKEN = os.environ["BOT_TOKEN"]  # set this in Railway's Variables tab — never hardcode it

# Portable temp dir: tempfile.gettempdir() respects $TMPDIR, so this resolves
# to a writable path on both Railway (/tmp) and Termux ($PREFIX/tmp) — a
# hardcoded "/tmp" fails on Android, which has no writable /tmp.
IMG_BASE_DIR = os.path.join(tempfile.gettempdir(), "quizician_imgs")

# ── Replace with YOUR Telegram numeric user ID ──────────────────
# To find it: message @userinfobot on Telegram → it replies with your ID
ADMIN_ID = 940770584

# ── Replace with your private GROUP's chat ID ────────────────────
# 1. Create the group, add this bot to it as a member (admin not required
#    unless you want it to survive being demoted/re-added later).
# 2. Send any message in the group, then send /storage_id in the SAME
#    group — the bot will reply with the chat ID (a negative number,
#    e.g. -1001234567890). Paste it below.
STORAGE_GROUP_ID = -1004447646576

# ── Replace with your quiz CHANNEL's chat ID ─────────────────────
# 1. Create a channel, add this bot as an ADMIN (channels require admin
#    rights for the bot to receive posts at all).
# 2. Forward any message from that channel to the bot in a private DM,
#    then send /quiz_channel_id right after — the bot replies with the ID.
QUIZ_CHANNEL_ID = -1004402622263

# ── Dedicated group for analytics JSON backups ────────────────
# The bot pins the latest analytics.json here after every change
# and deletes the previous one — always exactly one file in the group.
ANALYTICS_GROUP_ID = -1003767364410

# ── Curriculum structure for the quiz channel ─────────────────────
# Add new modules/subjects here as they come up. Lecture titles posted in
# the quiz channel must be formatted as:
#   "<Module> - <Subject> Lecture <number>: <name>"
#   e.g. "Endocrine - Physio Lecture 3: Insulin Signaling"
# Matching against this dict is case-insensitive; the canonical spelling
# below is what gets stored/displayed.
MODULES = {
    "Endocrine":      ["Bio", "Physio", "Patho", "Histo", "Pharma"],
    "Genitourinary":  ["Anatomy", "Physio", "Histo", "Patho", "Micro"],
}

# ═══════════════════════════════════════════════════════════════
# QUIZZY — The Quizician's cat friend 🐾
# ═══════════════════════════════════════════════════════════════
QUIZZY_WELCOME_ART = (
    " /\\_/\\ \n"
    "( ⌒.⌒ )\n"
    "  > ^ <  "
)
QUIZZY_SLEEPING_ART = (
    " /\\_/\\ \n"
    "(  -.- ) zzz\n"
    " > ^ <  "
)
# No "oops" expression was provided yet — this one's improvised to match
# the same style. Swap QUIZZY_OOPS_ART for a real one whenever you draw it.
QUIZZY_OOPS_ART = (
    " /\\_/\\ \n"
    "( ×_× )\n"
    " > ~ <  "
)

QUIZZY_WELCOME_LINES = [
    "صباح (أو مساء) الورد 🌹",
    "باشا البلد",
    "الله أكبر أخيرا قررت تذاكر",
]
QUIZZY_SUCCESS_LINES = [
    "تحياتي 🫡",
    "مش بقول باشا 😎",
    "قدوة 😌🙌",
]
QUIZZY_ERROR_LINES = [
    "كويزي وقع على دماغه من الصدمة، بس متقلقش هنظبطها 😓",
    "كويزي شايف إن المشكلة دي معندهاش داعي، جرب تاني 😓",
    "احنا مش عارفين إيه اللي حصل، بس كويزي واثق إنها هتتحل 😓",
    "كله بسبب قسم الفسيو 😓",
]

def quizzy_block(art: str, line: str) -> str:
    """Quizzy's ASCII art + one of his lines, wrapped for Telegram HTML.
    The art contains literal < > characters (whiskers/paws) which Telegram's
    HTML parser would otherwise choke on as broken tags — escape them."""
    return f"<pre>{html.escape(art)}</pre>\n<i>{html.escape(line)}</i>"

# ═══════════════════════════════════════════════════════════════
# BOT MESSAGES — every user-facing string the bot sends, in one place.
# Grouped by feature. Dynamic ones use {placeholders} filled with .format().
# (Content generated in a loop — like /c's command list, /quiz_list's
# lecture rows — stays where it's built, since there's
# nothing fixed to centralize there; only their static labels live here.)
# ═══════════════════════════════════════════════════════════════

# ── Generic / shared ──────────────────────────────────────────
MSG_ADMIN_ONLY = "🚫 للأدمن فقط"

# ── PDF collection flow ───────────────────────────────────────
MSG_PDF_ASK_NAME = (
    "✏️ <b>اكتب اسم التوحفة الفنية (الملف) اللي عايزه:</b>\n"
    "<i>Lecture 1 Anatomy Questions</i>"
)
MSG_PDF_EMPTY = "❌ لا يوجد أسئلة محفوظة"
MSG_EXPORT_EMPTY = "❌ لا يوجد أسئلة محفوظة بعد"
MSG_EXPORT_GENERATING = "⏳ جاري توليد {kind} لـ {count} عنصر..."
MSG_PDF_GENERATING = "⏳ جاري توليد PDF لـ {count} عنصر..."
MSG_PDF_CAPTION = "📄 {count} سؤال — {name} ❤️\n\n <i>{quizzy_line}</i>"
MSG_DOCX_CAPTION = "📝 {count} سؤال — {name} ❤️\n\n <i>{quizzy_line}</i>"
MSG_DOCX_UNAVAILABLE = (
    "❌ DOCX export مش متاح دلوقتي (python-docx مش متثبت). "
    "استخدم PDF Export بدل كده، أو ثبّت python-docx وأعد التشغيل."
)
MSG_PDF_CLEARED = "🗑 تم قرار إزالة يا دولي"
MSG_EXPORT_CLEARED_ALL = "🗑 تم قرار إزاله يا دولي"
MSG_CANCEL_DONE = "❌ تم نطر أبلكاش"
MSG_CANCEL_NOTHING = "بتلغيني أنا يعني ولا أي🤨"

# ═══════════════════════════════════════════════════════════════
# USERS STORAGE
# ═══════════════════════════════════════════════════════════════
USERS_FILE = "users.json"

def load_users():
    if os.path.exists(USERS_FILE):
        with open(USERS_FILE, "r") as f:
            return set(json.load(f))
    return set()

def save_users():
    with open(USERS_FILE, "w") as f:
        json.dump(list(USERS), f)

USERS = load_users()

# ═══════════════════════════════════════════════════════════════
# ANALYTICS — XP · LEVELS · ACHIEVEMENTS
#
# analytics.json schema per user:
# {
#   "questions_created": int,
#   "streak":            int,
#   "last_active_date":  "YYYY-MM-DD" | null,
#   "pdfs_exported":     int,
#   "xp":                int,
#   "level":             int,
#   "achievements": {
#       "questions": 0-5,   # tiers unlocked
#       "streak":    0-5,
#       "pdfs":      0-5,
#       "speed":     0-5
#   }
# }
# ═══════════════════════════════════════════════════════════════
ANALYTICS_FILE          = "analytics.json"
ANALYTICS_BACKUP_MARKER = "🗄 QUIZICIAN_ANALYTICS_BACKUP"


# ── XP per action ────────────────────────────────────────────
XP_PER_QUESTION   = 10
XP_PDF_EXPORT     = 25   # flat bonus on top of per-question XP

# ── Level curve: level = floor(sqrt(xp / XP_LEVEL_FACTOR)) ──
XP_LEVEL_FACTOR   = 50   # level 1 = 50xp, 5 = 1250xp, 10 = 5000xp

# ── Achievement definitions ───────────────────────────────────
# Each achievement has 5 tiers: (threshold, label, xp_bonus, emoji)
ACHIEVEMENTS = {
    "questions": [
        (1,    "صانع الأسئلة I",    25,  "❓"),
        (100,  "صانع الأسئلة II",   75,  "❓"),
        (250,  "صانع الأسئلة III", 150,  "❓"),
        (500,  "صانع الأسئلة IV",  300,  "❓"),
        (1000, "صانع الأسئلة V",   600,  "❓"),
    ],
    "streak": [
        (3,   "ملتزم I",    20,  "🔥"),
        (7,   "ملتزم II",   50,  "🔥"),
        (30,  "ملتزم III", 150,  "🔥"),
        (100, "ملتزم IV",  400,  "🔥"),
        (365, "ملتزم V",   1000, "🔥"),
    ],
    "pdfs": [
        (1,  "صانع PDF I",    30,  "📚"),
        (5,  "صانع PDF II",   80,  "📚"),
        (10, "صانع PDF III", 175,  "📚"),
        (25, "صانع PDF IV",  400,  "📚"),
        (50, "صانع PDF V",   800,  "📚"),
    ],
    "speed": [
        (5,  "سريع I",    20,  "⚡"),
        (10, "سريع II",   50,  "⚡"),
        (20, "سريع III", 125,  "⚡"),
        (30, "سريع IV",  250,  "⚡"),
        (50, "سريع V",   500,  "⚡"),
    ],
}

LEVEL_TITLES = {
    0:  "مبتدئ",
    1:  "متعلم",
    3:  "نشيط",
    5:  "محترف",
    8:  "خبير",
    12: "أستاذ",
    17: "أسطورة",
}

def _level_title(level: int) -> str:
    title = LEVEL_TITLES[0]
    for threshold, t in LEVEL_TITLES.items():
        if level >= threshold:
            title = t
    return title

def _xp_to_level(xp: int) -> int:
    import math
    return int(math.floor(math.sqrt(xp / XP_LEVEL_FACTOR)))

def _level_xp_range(level: int) -> tuple[int, int]:
    """(xp_start, xp_end) for this level."""
    return level ** 2 * XP_LEVEL_FACTOR, (level + 1) ** 2 * XP_LEVEL_FACTOR

def _blank_entry() -> dict:
    return {
        "questions_created": 0,
        "streak":            0,
        "last_active_date":  None,
        "pdfs_exported":     0,
        "xp":                0,
        "level":             0,
        "achievements":      {k: 0 for k in ACHIEVEMENTS},
    }

def load_analytics() -> dict:
    if os.path.exists(ANALYTICS_FILE):
        with open(ANALYTICS_FILE) as f:
            return json.load(f)
    return {}

def save_analytics():
    with open(ANALYTICS_FILE, "w") as f:
        json.dump(ANALYTICS, f, indent=2)

ANALYTICS: dict = load_analytics()

# Message ID of the currently pinned analytics backup in ANALYTICS_GROUP_ID.
# Populated on startup by restore_analytics_from_channel; the pin is the
# source of truth — no separate state file needed.
_analytics_backup_msg_id: int | None = None

def _today() -> str:
    from datetime import datetime, timezone
    return datetime.now(timezone.utc).strftime("%Y-%m-%d")

def _get_entry(user_id: int) -> dict:
    key   = str(user_id)
    entry = ANALYTICS.setdefault(key, _blank_entry())
    # backfill missing keys for users created before this system
    for k, v in _blank_entry().items():
        entry.setdefault(k, v)
    if not isinstance(entry.get("achievements"), dict):
        entry["achievements"] = {k: 0 for k in ACHIEVEMENTS}
    for k in ACHIEVEMENTS:
        entry["achievements"].setdefault(k, 0)
    return entry

def _award_xp(entry: dict, amount: int) -> int:
    """Add XP, recalculate level. Returns new level if levelled up, else 0."""
    entry["xp"] += amount
    new_level    = _xp_to_level(entry["xp"])
    levelled_up  = new_level > entry["level"]
    entry["level"] = new_level
    return new_level if levelled_up else 0

def _check_achievements(entry: dict, stat_key: str) -> list[dict]:
    """Check one stat against its achievement tiers. Returns list of newly
    unlocked tiers as dicts with keys: name, emoji, xp_bonus, tier (1-5)."""
    value    = entry.get(stat_key, 0)
    tiers    = ACHIEVEMENTS[stat_key]
    current  = entry["achievements"][stat_key]
    unlocked = []
    for i, (threshold, name, xp_bonus, emoji) in enumerate(tiers):
        tier = i + 1
        if tier <= current:
            continue
        if value >= threshold:
            entry["achievements"][stat_key] = tier
            unlocked.append({"name": name, "emoji": emoji,
                              "xp_bonus": xp_bonus, "tier": tier})
            _award_xp(entry, xp_bonus)
        else:
            break   # tiers are ordered — no point checking higher ones
    return unlocked

def _record_activity(user_id: int, questions_delta: int = 0,
                     pdfs_delta: int = 0, session_questions: int = 0) -> dict:
    """Update all stats. Returns dict of events for the caller to announce:
    { "achievements": [...], "level_up": int | 0 }"""
    from datetime import datetime, timezone, timedelta
    today  = _today()
    entry  = _get_entry(user_id)
    last   = entry.get("last_active_date")

    # ── streak ───────────────────────────────────────────────
    if last != today:
        yesterday = (datetime.now(timezone.utc) - timedelta(days=1)).strftime("%Y-%m-%d")
        entry["streak"] = (entry["streak"] + 1) if last == yesterday else 1
        entry["last_active_date"] = today

    # ── counters ─────────────────────────────────────────────
    entry["questions_created"] += questions_delta
    entry["pdfs_exported"]     += pdfs_delta

    # ── XP for raw actions ───────────────────────────────────
    xp_earned = questions_delta * XP_PER_QUESTION + pdfs_delta * XP_PDF_EXPORT
    new_level  = _award_xp(entry, xp_earned) if xp_earned else 0

    # ── achievement checks ───────────────────────────────────
    newly_unlocked = []
    newly_unlocked += _check_achievements(entry, "questions")
    newly_unlocked += _check_achievements(entry, "streak")
    newly_unlocked += _check_achievements(entry, "pdfs")

    # speed: session_questions is how many were made in one PDF session
    if session_questions:
        # temporarily set a "session_max" so we can check the speed tier
        # using the same threshold logic; we don't persist session count
        prev = entry.get("_session_max", 0)
        if session_questions > prev:
            entry["_session_max"] = session_questions
            newly_unlocked += _check_achievements(entry, "speed")

    # level-up might also happen from achievement XP bonuses
    final_level = _xp_to_level(entry["xp"])
    if final_level > entry["level"]:
        entry["level"] = final_level
        new_level = final_level

    save_analytics()
    return {"achievements": newly_unlocked, "level_up": new_level}

async def _announce_events(context, chat_id: int, events: dict):
    """Send achievement unlocks and level-up notifications to the user's DM."""
    msgs = []

    for ach in events.get("achievements", []):
        stars = "⭐" * ach["tier"]
        msgs.append(
            f"{ach['emoji']} <b>إنجاز جديد!</b>\n"
            f"<b>{ach['name']}</b> {stars}\n"
            f"<i>+{ach['xp_bonus']} XP</i>"
        )

    if events.get("level_up"):
        lvl   = events["level_up"]
        title = _level_title(lvl)
        msgs.append(
            f"🎉 <b>ترقية!</b>\n"
            f"وصلت للمستوى <b>{lvl}</b> — <i>{title}</i>"
        )

    for msg in msgs:
        await context.bot.send_message(
            chat_id=chat_id, text=msg, parse_mode=ParseMode.HTML
        )

async def backup_analytics_to_channel(context):
    global _analytics_backup_msg_id
    if not ANALYTICS_GROUP_ID:
        return
    data = json.dumps(ANALYTICS, indent=2).encode("utf-8")
    try:
        sent = await context.bot.send_document(
            chat_id=ANALYTICS_GROUP_ID,
            document=InputFile(BytesIO(data), filename="analytics.json"),
            caption=ANALYTICS_BACKUP_MARKER,
        )
    except Exception as e:
        print("ANALYTICS BACKUP ERROR:", e)
        return
    try:
        await context.bot.pin_chat_message(
            chat_id=ANALYTICS_GROUP_ID,
            message_id=sent.message_id,
            disable_notification=True,
        )
    except Exception as e:
        print("ANALYTICS PIN ERROR:", e)
    if _analytics_backup_msg_id and _analytics_backup_msg_id != sent.message_id:
        try:
            await context.bot.delete_message(
                chat_id=ANALYTICS_GROUP_ID,
                message_id=_analytics_backup_msg_id,
            )
        except Exception:
            pass
    _analytics_backup_msg_id = sent.message_id

async def restore_analytics_from_channel(app):
    global _analytics_backup_msg_id
    if not ANALYTICS_GROUP_ID:
        return
    try:
        chat   = await app.bot.get_chat(ANALYTICS_GROUP_ID)
        pinned = chat.pinned_message
        if not pinned or not pinned.document:
            return
        if (pinned.caption or "") != ANALYTICS_BACKUP_MARKER:
            return
        tg_file = await app.bot.get_file(pinned.document.file_id)
        raw     = await tg_file.download_as_bytearray()
        ANALYTICS.update(json.loads(bytes(raw).decode("utf-8")))
        save_analytics()
        _analytics_backup_msg_id = pinned.message_id
        print(f"Restored analytics: {len(ANALYTICS)} user(s).")
    except Exception as e:
        print("ANALYTICS RESTORE ERROR:", e)

# ═══════════════════════════════════════════════════════════════
# PASSWORD-GATED STORAGE (private group)
# ═══════════════════════════════════════════════════════════════
# STORAGE_GROUP_ID is the vault: post any photo/video/document/album there
# with a caption starting with a password word, and the bot indexes it.
# A DM containing that exact word gets the item(s) copied to the user.
STORAGE_INDEX_FILE = "storage_index.json"

def load_storage_index():
    if os.path.exists(STORAGE_INDEX_FILE):
        with open(STORAGE_INDEX_FILE, "r") as f:
            return json.load(f)
    return {}

def save_storage_index():
    with open(STORAGE_INDEX_FILE, "w") as f:
        json.dump(STORAGE_INDEX, f)

# password (lowercased) -> list of items; each item is a list of message_ids
# (a single-message item is [id], an album is [id1, id2, ...]). Reusing the
# same password just appends another item — both get delivered on unlock.
STORAGE_INDEX: dict = load_storage_index()

# media_group_id -> {"ids": [...], "caption": str|None, "task": asyncio.Task}
# Albums arrive as several separate updates; we debounce them so the whole
# album gets filed as one item under one password.
ALBUM_BUFFER: dict = {}

# ── Durable backup: the local JSON files above are only a same-host cache.
# A bot can't scan a channel's history, but it CAN always read a chat's
# currently pinned message on demand — restart, redeploy, or host switch
# doesn't matter. So we mirror USERS + STORAGE_INDEX into one pinned
# message in the storage group itself, and rebuild the local cache from
# it on startup if the local files are ever missing/wiped.
STORAGE_BACKUP_MARKER     = "🗄 QUIZICIAN_STORAGE_BACKUP"
STORAGE_BACKUP_STATE_FILE = "storage_backup_state.json"

def load_storage_backup_state():
    if os.path.exists(STORAGE_BACKUP_STATE_FILE):
        with open(STORAGE_BACKUP_STATE_FILE, "r") as f:
            return json.load(f)
    return {}

def save_storage_backup_state():
    with open(STORAGE_BACKUP_STATE_FILE, "w") as f:
        json.dump(STORAGE_BACKUP_STATE, f)

STORAGE_BACKUP_STATE: dict = load_storage_backup_state()  # {"backup_msg_id": int}

async def backup_storage_to_channel(context: ContextTypes.DEFAULT_TYPE):
    if not STORAGE_GROUP_ID:
        return
    payload = {"users": list(USERS), "storage_index": STORAGE_INDEX}
    data    = json.dumps(payload).encode("utf-8")
    # A pinned document instead of a pinned text message: Bot API caps
    # documents at 50MB vs ~4KB for a text message — effectively removes
    # the size ceiling for any realistic amount of data this bot handles.
    filename = "quizician_storage_backup.json"

    # We tried editing the existing pinned document in place, but Telegram
    # reliably rejected it with "Can't parse inputmedia: media not found".
    # Simpler and just as effective: send a new document, pin it, then
    # delete the previous one — net result is still exactly one backup
    # document sitting in the chat at all times.
    old_msg_id = STORAGE_BACKUP_STATE.get("backup_msg_id")

    try:
        sent = await context.bot.send_document(
            chat_id=STORAGE_GROUP_ID,
            document=InputFile(BytesIO(data), filename=filename),
            caption=STORAGE_BACKUP_MARKER,
        )
    except Exception as e:
        print("STORAGE BACKUP ERROR:", e)
        return

    # Save the message id immediately — pinning/deleting-old are nice-to-
    # haves on top, and their failure (e.g. bot isn't admin / lacks rights)
    # must NOT stop us from remembering this new message id.
    STORAGE_BACKUP_STATE["backup_msg_id"] = sent.message_id
    save_storage_backup_state()

    try:
        await context.bot.pin_chat_message(chat_id=STORAGE_GROUP_ID, message_id=sent.message_id, disable_notification=True)
    except Exception as e:
        print("STORAGE BACKUP PIN ERROR (message saved anyway, but won't be pinned — check bot is admin with pin rights):", e)

    if old_msg_id and old_msg_id != sent.message_id:
        try:
            await context.bot.delete_message(chat_id=STORAGE_GROUP_ID, message_id=old_msg_id)
        except Exception as e:
            print("STORAGE BACKUP OLD-MESSAGE DELETE ERROR (probably already gone, harmless):", e)

async def restore_storage_from_channel(app):
    """Runs once on startup — rebuilds USERS + STORAGE_INDEX from the
    storage group's pinned backup if the local cache is missing/stale."""
    if not STORAGE_GROUP_ID:
        return
    try:
        chat   = await app.bot.get_chat(STORAGE_GROUP_ID)
        pinned = chat.pinned_message
        if pinned and pinned.document and (pinned.caption or "") == STORAGE_BACKUP_MARKER:
            tg_file = await app.bot.get_file(pinned.document.file_id)
            raw     = await tg_file.download_as_bytearray()
            payload = json.loads(bytes(raw).decode("utf-8"))
            USERS.update(payload.get("users", []))
            STORAGE_INDEX.update(payload.get("storage_index", {}))
            save_users()
            save_storage_index()
            STORAGE_BACKUP_STATE["backup_msg_id"] = pinned.message_id
            save_storage_backup_state()
            print(f"Restored storage backup: {len(USERS)} user(s), {len(STORAGE_INDEX)} password(s).")
    except Exception as e:
        if "chat not found" in str(e).lower():
            print(
                "STORAGE RESTORE ERROR: Chat not found — STORAGE_GROUP_ID "
                f"({STORAGE_GROUP_ID}) isn't a real group this bot knows about. "
                "Still the template placeholder, wrong ID, or the bot was never "
                "added to that group. See the setup comment above STORAGE_GROUP_ID."
            )
        else:
            print("STORAGE RESTORE ERROR:", e)

# ═══════════════════════════════════════════════════════════════
# QUIZ CHANNEL (interactive quiz storage, organized by lecture)
# ═══════════════════════════════════════════════════════════════
# Post plain text in QUIZ_CHANNEL_ID to open/resume a lecture (that text
# becomes the lecture's name), then post quiz polls one by one — each gets
# filed under the currently-open lecture, in posting order. Post "-END" to
# close the lecture. Users pick a closed lecture via /quiz and the bot
# delivers the ORIGINAL polls via copy_messages (fresh, unattributed,
# independently answerable — no send_poll(), no need to know the answer).
QUIZ_INDEX_FILE = "quiz_index.json"

def load_quiz_index():
    if os.path.exists(QUIZ_INDEX_FILE):
        with open(QUIZ_INDEX_FILE, "r") as f:
            return json.load(f)
    return {}

def save_quiz_index():
    with open(QUIZ_INDEX_FILE, "w") as f:
        json.dump(QUIZ_INDEX, f)

QUIZ_INDEX: dict = load_quiz_index()  # lecture_name -> {"ids": [...], "closed": bool, "module": str, "subject": str, "lecture_number": str, "name": str}

# Matches "<Subject> Lecture <number>", e.g. "Physio Lecture 3"
_LECTURE_TITLE_RE = re.compile(r"^(.*?)\s+Lecture\s+(\d+)\s*$", re.IGNORECASE)

def parse_lecture_title(text: str):
    """Parses "<Module> - <Subject> Lecture <number>: <name>" against the
    MODULES curriculum. Returns (module, subject, lecture_number, name) on
    success, or (None, None, None, error_message) on failure — matching is
    case-insensitive but the canonical spelling from MODULES is returned."""
    if " - " not in text or ":" not in text:
        return None, None, None, (
            "⚠️ الصيغة غلط. لازم تكون:\n"
            "<code>Module - Subject Lecture Number: Name</code>\n"
            "مثال: <code>Endocrine - Physio Lecture 3: Insulin Signaling</code>"
        )
    module_part, rest = text.split(" - ", 1)
    subj_lec_part, name = rest.split(":", 1)
    module_part, subj_lec_part, name = module_part.strip(), subj_lec_part.strip(), name.strip()

    module_match = next((m for m in MODULES if m.lower() == module_part.lower()), None)
    if not module_match:
        valid = ", ".join(MODULES.keys())
        return None, None, None, f"⚠️ الموديول \"{module_part}\" مش معروف. الموديولات المتاحة: {valid}"

    m = _LECTURE_TITLE_RE.match(subj_lec_part)
    if not m:
        return None, None, None, (
            "⚠️ الصيغة غلط بعد اسم الموديول. لازم تكون:\n"
            "<code>Subject Lecture Number</code>\n"
            "مثال: <code>Physio Lecture 3</code>"
        )
    subject_part, lecture_number = m.group(1).strip(), m.group(2).strip()
    subject_match = next((s for s in MODULES[module_match] if s.lower() == subject_part.lower()), None)
    if not subject_match:
        valid = ", ".join(MODULES[module_match])
        return None, None, None, f"⚠️ المادة \"{subject_part}\" مش من موديول {module_match}. المواد المتاحة: {valid}"

    return module_match, subject_match, lecture_number, name

def ready_modules():
    # Always show every configured module — even ones with zero lectures
    # posted yet — so the curriculum structure is visible from day one.
    return list(MODULES.keys())

def ready_subjects(module: str):
    # Same idea: every subject defined for this module shows up, regardless
    # of whether any lecture has been posted for it yet.
    return list(MODULES.get(module, []))

def ready_lecture_keys(module: str, subject: str):
    return [
        name for name, v in QUIZ_INDEX.items()
        if v["closed"] and v["ids"] and v["module"] == module and v["subject"] == subject
    ]  # insertion order = numbering order

QUIZ_STATE_FILE = "quiz_state.json"

def load_quiz_state():
    if os.path.exists(QUIZ_STATE_FILE):
        with open(QUIZ_STATE_FILE, "r") as f:
            return json.load(f)
    return {"current_lecture": None}

def save_quiz_state():
    with open(QUIZ_STATE_FILE, "w") as f:
        json.dump(QUIZ_STATE, f)

QUIZ_STATE: dict = load_quiz_state()  # survives restarts mid-lecture

QUIZ_POLL_STATUS_FILE = "quiz_poll_status.json"

def load_quiz_poll_status():
    if os.path.exists(QUIZ_POLL_STATUS_FILE):
        with open(QUIZ_POLL_STATUS_FILE, "r") as f:
            return json.load(f)
    return {}

def save_quiz_poll_status():
    with open(QUIZ_POLL_STATUS_FILE, "w") as f:
        json.dump(QUIZ_POLL_STATUS, f)

# poll_id -> {"lecture": str, "message_id": int, "closed": bool}
# Tracks whether each quiz-channel poll has been stopped yet — Telegram
# only allows copying a quiz poll once its correct answer is known, i.e.
# once it's been stopped, so this is what /quiz delivery checks against.
QUIZ_POLL_STATUS: dict = load_quiz_poll_status()

# ── Durable backup: same pinned-message trick as the storage group, so
# the lecture/quiz index survives a host switch or wiped local disk —
# only the local JSON cache is fragile, the channel content itself never
# was.
QUIZ_BACKUP_MARKER     = "🗄 QUIZICIAN_QUIZ_BACKUP"
QUIZ_BACKUP_STATE_FILE = "quiz_backup_state.json"

def load_quiz_backup_state():
    if os.path.exists(QUIZ_BACKUP_STATE_FILE):
        with open(QUIZ_BACKUP_STATE_FILE, "r") as f:
            return json.load(f)
    return {}

def save_quiz_backup_state():
    with open(QUIZ_BACKUP_STATE_FILE, "w") as f:
        json.dump(QUIZ_BACKUP_STATE, f)

QUIZ_BACKUP_STATE: dict = load_quiz_backup_state()  # {"backup_msg_id": int}

async def backup_quiz_to_channel(context: ContextTypes.DEFAULT_TYPE):
    if not QUIZ_CHANNEL_ID:
        return
    payload  = {"quiz_index": QUIZ_INDEX, "quiz_state": QUIZ_STATE, "quiz_poll_status": QUIZ_POLL_STATUS}
    data     = json.dumps(payload).encode("utf-8")
    filename = "quizician_quiz_backup.json"

    # Same approach as storage backup: editing the existing pinned document
    # in place reliably failed with "Can't parse inputmedia: media not
    # found", so instead we send a new one, pin it, then delete the old —
    # still exactly one backup document in the channel at any time.
    old_msg_id = QUIZ_BACKUP_STATE.get("backup_msg_id")

    try:
        sent = await context.bot.send_document(
            chat_id=QUIZ_CHANNEL_ID,
            document=InputFile(BytesIO(data), filename=filename),
            caption=QUIZ_BACKUP_MARKER,
        )
    except Exception as e:
        print("QUIZ BACKUP ERROR:", e)
        return

    # Same fix as storage backup: persist the message id regardless of
    # whether pinning/deleting-old succeed, so we don't resend a fresh
    # document on every single addition.
    QUIZ_BACKUP_STATE["backup_msg_id"] = sent.message_id
    save_quiz_backup_state()

    try:
        await context.bot.pin_chat_message(chat_id=QUIZ_CHANNEL_ID, message_id=sent.message_id, disable_notification=True)
    except Exception as e:
        print("QUIZ BACKUP PIN ERROR (message saved anyway, but won't be pinned — check bot is admin with pin rights):", e)

    if old_msg_id and old_msg_id != sent.message_id:
        try:
            await context.bot.delete_message(chat_id=QUIZ_CHANNEL_ID, message_id=old_msg_id)
        except Exception as e:
            print("QUIZ BACKUP OLD-MESSAGE DELETE ERROR (probably already gone, harmless):", e)

async def restore_quiz_from_channel(app):
    """Runs once on startup — rebuilds the lecture/quiz index from the quiz
    channel's pinned backup if the local cache is missing/stale."""
    if not QUIZ_CHANNEL_ID:
        return
    try:
        chat   = await app.bot.get_chat(QUIZ_CHANNEL_ID)
        pinned = chat.pinned_message
        if pinned and pinned.document and (pinned.caption or "") == QUIZ_BACKUP_MARKER:
            tg_file = await app.bot.get_file(pinned.document.file_id)
            raw     = await tg_file.download_as_bytearray()
            payload = json.loads(bytes(raw).decode("utf-8"))
            QUIZ_INDEX.update(payload.get("quiz_index", {}))
            QUIZ_STATE.update(payload.get("quiz_state", {}))
            QUIZ_POLL_STATUS.update(payload.get("quiz_poll_status", {}))
            save_quiz_index()
            save_quiz_state()
            save_quiz_poll_status()
            QUIZ_BACKUP_STATE["backup_msg_id"] = pinned.message_id
            save_quiz_backup_state()
            print(f"Restored quiz backup: {len(QUIZ_INDEX)} lecture(s).")
    except Exception as e:
        if "chat not found" in str(e).lower():
            print(
                "QUIZ RESTORE ERROR: Chat not found — QUIZ_CHANNEL_ID "
                f"({QUIZ_CHANNEL_ID}) isn't a real channel this bot knows about. "
                "Still the template placeholder, wrong ID, or the bot was never "
                "added as an admin there. See the setup comment above QUIZ_CHANNEL_ID."
            )
        else:
            print("QUIZ RESTORE ERROR:", e)

# ═══════════════════════════════════════════════════════════════
# STATE
# ═══════════════════════════════════════════════════════════════
PDF_BUFFER             = {}    # user_id -> list of item dicts
PDF_NAMES              = {}    # user_id -> str
AWAITING_NAME          = {}    # user_id -> True
SLEEPING               = set()
PROGRESS_MSG_ID        = {}    # user_id -> message_id of the live progress message
PENDING_IMAGE          = {}    # user_id -> local path of an image awaiting its question
CLARIFY_QUEUE          = {}    # user_id -> list of PDF_BUFFER indices awaiting a correct-answer tap
POLL_WATCH             = {}    # poll_id -> (user_id, item_index) for passive auto-detection
PENDING_EDIT           = {}    # user_id -> {"index": int, "field": "q"/"title"/"content"/"option", "opt_index": int?}
                                # awaiting free-text replacement for one field of a just-added question
LECTURE_SESSIONS       = {}    # user_id -> {"module","subject","lecture_key","queue":[mid,...],
                                #             "current_poll_id","total","answered"} — active one-at-a-time delivery

# ═══════════════════════════════════════════════════════════════
# CONSTANTS
# ═══════════════════════════════════════════════════════════════
MAX_QUESTIONS_PER_MSG = 40
TELEGRAM_Q_LIMIT      = 300   # max chars in poll question field
TELEGRAM_DESC_LIMIT   = 200   # max chars in poll description (shown above question)
TELEGRAM_EX_LIMIT     = 200   # max chars in poll explanation (shown after answering)
PDF_MAX_IMG_WIDTH     = 13 * cm

# ═══════════════════════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════════════════════
def clean_option(line: str) -> str:
    line = line.strip()
    line = re.sub(r"^[A-Ea-e1-5][).\-]\s*", "", line)
    line = re.sub(r"^[-•]\s*", "", line)
    return line.strip()

def strip_leading_letter_prefix(option: str) -> str:
    return re.sub(r"^[A-Ea-e]\)\s*", "", option).strip()

def normalize_mcq_block(block: str):
    block = block.strip()
    if "\n" in block:
        return [l.strip() for l in block.split("\n") if l.strip()]
    match = re.search(r"\b([A-Ea-e1-5])[).]", block)
    if not match:
        return [block]
    question     = block[:match.start()].strip()
    options_part = block[match.start():]
    parts = re.split(r"(?=\b[A-Ea-e1-5][).])", options_part)
    return [question] + [p.strip() for p in parts if p.strip()]

def strip_spoiler_markers(text: str) -> str:
    return re.sub(r"\|\|(.+?)\|\|", r"\1", text, flags=re.DOTALL)

def parse_mcq_lines(lines: list):
    """
    Given already-normalized MCQ lines (question line + option lines),
    extracts (question, raw_options, correct_index, explanation).
    correct_index is None if no option was marked correct.
    Shared by the text handler and the image-caption parser so the
    MCQ grammar only lives in one place.
    """
    question      = lines[0]
    options       = []
    correct_index = None
    explanation   = None

    for line in lines[1:]:
        ex_match = re.match(r"^ex:\s*(.+)", line, re.IGNORECASE)
        if ex_match:
            explanation = ex_match.group(1).strip()
            continue

        opt       = clean_option(line)
        has_z_end = re.search(r"\s+[zZ]\s*$", opt)
        has_check = "✅" in opt

        if has_z_end or has_check:
            opt           = opt.replace("✅", "")
            opt           = re.sub(r"\s+[zZ]\s*$", "", opt).strip()
            correct_index = len(options)

        if opt:
            options.append(opt)

    return question, options, correct_index, explanation

def parse_mcq_block(block: str):
    """
    Full validation on top of parse_mcq_lines: returns
    (question, raw_options, correct_index, explanation) only if the block
    is a COMPLETE, valid MCQ (>=3 lines, a correct answer marked).
    Returns None otherwise. Used to detect a fully-formed question in an
    image caption so we don't need to ask the user to resend it.
    """
    lines = normalize_mcq_block(block.strip())
    if len(lines) < 3:
        return None
    question, options, correct_index, explanation = parse_mcq_lines(lines)
    if correct_index is None or correct_index >= len(options):
        return None
    return question, options, correct_index, explanation

def parse_written_question(block: str):
    block = strip_spoiler_markers(block)
    lines = [l.rstrip() for l in block.split("\n") if l.strip()]
    if len(lines) < 2:
        return None
    title = re.sub(r'[\""\']+$', "", lines[0]).strip()
    content_lines = lines[1:]
    content = "\n".join(content_lines).strip()
    if not content:
        return None
    if content.startswith(".") and content.endswith("."):
        content = content[1:-1].strip()
        return title, content
    if content:
        return title, content
    return None

def parse_written_strict(block: str):
    block = strip_spoiler_markers(block)
    lines = [l.rstrip() for l in block.split("\n") if l.strip()]
    if len(lines) < 2:
        return None
    title   = lines[0]
    content = "\n".join(lines[1:]).strip()
    if content.startswith(".") and content.endswith("."):
        return title, content[1:-1].strip()
    return None

def split_question_for_telegram(question: str):
    """
    Returns (main_q, description_overflow) where:
    - main_q fits in TELEGRAM_Q_LIMIT
    - description_overflow goes into the 'description' field (shown above question)
      and is capped at TELEGRAM_DESC_LIMIT
    If question fits in Q_LIMIT, description_overflow is None.
    """
    if len(question) <= TELEGRAM_Q_LIMIT:
        return question, None
    # Try to split at a sentence boundary
    cutoff    = TELEGRAM_Q_LIMIT - 3
    split_pos = question.rfind(". ", 0, cutoff)
    if split_pos == -1:
        split_pos = question.rfind(" ", 0, cutoff)
    if split_pos == -1:
        split_pos = cutoff
    main     = question[:split_pos].strip() + "…"
    overflow = "…" + question[split_pos:].strip()
    # Cap overflow to TELEGRAM_DESC_LIMIT
    if len(overflow) > TELEGRAM_DESC_LIMIT:
        overflow = overflow[:TELEGRAM_DESC_LIMIT - 1] + "…"
    return main, overflow

def options_too_long(options: list) -> bool:
    """Check if any single option exceeds Telegram's 100-char option limit."""
    return any(len(o) > 100 for o in options)

def make_letter_only_options(count: int) -> list:
    """Return ['A', 'B', 'C', ...] for poll when answers are too long."""
    return [string.ascii_uppercase[i] for i in range(count)]

def _cleanup_images(user_id: int):
    import shutil
    img_dir = os.path.join(IMG_BASE_DIR, str(user_id))
    if os.path.exists(img_dir):
        shutil.rmtree(img_dir, ignore_errors=True)

def _clear_pending_image(user_id: int):
    """Drop any image that's still waiting for a question, deleting its file."""
    path = PENDING_IMAGE.pop(user_id, None)
    if path and os.path.exists(path):
        try:
            os.remove(path)
        except Exception:
            pass

def _clear_pending_edit(user_id: int):
    """Drop any pending 'send new text for this field' state for this user."""
    PENDING_EDIT.pop(user_id, None)

def _clear_clarify_queue(user_id: int):
    """Drop any pending 'choose the correct answer' queue/watches for this user."""
    CLARIFY_QUEUE.pop(user_id, None)
    stale_poll_ids = [pid for pid, (uid, _) in POLL_WATCH.items() if uid == user_id]
    for pid in stale_poll_ids:
        POLL_WATCH.pop(pid, None)

# ═══════════════════════════════════════════════════════════════
# QUIZ DELIVERY  (single source of truth for sending a live quiz poll)
# ═══════════════════════════════════════════════════════════════
async def _send_quiz_poll(context, poll_kwargs: dict, image_path: str = None):
    """
    Sends the poll, attaching image_path as the quiz's native media (Bot API
    10.0+ InputPollMedia) when provided. Falls back to sending the image as a
    separate message + a media-less poll if the media attachment is ever
    rejected — this feature is new enough (May 2026) that we don't want a
    server-side quirk to silently drop the question entirely.
    """
    if image_path:
        try:
            with open(image_path, "rb") as f:
                await context.bot.send_poll(**poll_kwargs, media=InputMediaPhoto(f))
            return
        except Exception as e:
            print("POLL MEDIA ERROR (falling back to separate image message):", e)
            with open(image_path, "rb") as f:
                await context.bot.send_photo(chat_id=poll_kwargs["chat_id"], photo=f)
    await context.bot.send_poll(**poll_kwargs)

async def deliver_quiz(
    context, chat_id: int, question: str, raw_options: list, correct_index: int,
    explanation: str = None, image_path: str = None,
    always_show_question_text: bool = False, header_label: str = "📋 <b>السؤال:</b>",
):
    """
    Sends a single live quiz poll to chat_id, handling Telegram's field-length
    limits consistently (question <=300, options <=100, explanation <=200).
    If image_path is given, it's attached as the quiz's native photo
    attachment (Bot API 10.0+), so it shows up inside the quiz itself.

    always_show_question_text=True forces the original question text to be
    shown as a message even when it fits inside the poll's question field —
    used for forwarded quizzes so the original wording is always visible.

    This is the ONLY place that builds/sends quiz polls in non-PDF mode, so
    forwarded polls, typed MCQs, and image-paired MCQs all share one code path.
    """
    labeled_options = [
        f"{string.ascii_uppercase[i]}) {opt}" for i, opt in enumerate(raw_options)
    ]
    q_fits      = len(question) <= TELEGRAM_Q_LIMIT
    answers_fit = not options_too_long(labeled_options)

    if q_fits and answers_fit:
        main_q, desc_overflow = split_question_for_telegram(question)

        if always_show_question_text:
            await context.bot.send_message(
                chat_id=chat_id, text=f"{header_label}\n{question}", parse_mode=ParseMode.HTML,
            )

        if desc_overflow:
            await context.bot.send_message(
                chat_id=chat_id,
                text=f"📋 <b>تكملة السؤال:</b>\n{desc_overflow}",
                parse_mode=ParseMode.HTML,
            )

        poll_kwargs = dict(
            chat_id=chat_id, question=main_q, options=labeled_options,
            type="quiz", correct_option_id=correct_index, is_anonymous=True,
        )
        if explanation:
            poll_kwargs["explanation"] = explanation[:TELEGRAM_EX_LIMIT]
        await _send_quiz_poll(context, poll_kwargs, image_path)

    elif not q_fits and answers_fit:
        await context.bot.send_message(
            chat_id=chat_id, text=f"{header_label}\n{question}", parse_mode=ParseMode.HTML,
        )

        poll_kwargs = dict(
            chat_id=chat_id, question=".", options=labeled_options,
            type="quiz", correct_option_id=correct_index, is_anonymous=True,
        )
        if explanation:
            poll_kwargs["explanation"] = explanation[:TELEGRAM_EX_LIMIT]
        await _send_quiz_poll(context, poll_kwargs, image_path)

    else:
        answer_lines = "\n".join(
            f"{'✅ ' if i == correct_index else ''}{string.ascii_uppercase[i]}) {opt}"
            for i, opt in enumerate(raw_options)
        )
        await context.bot.send_message(
            chat_id=chat_id,
            text=f"{header_label}\n{question}\n\n<b>الإجابات:</b>\n{answer_lines}",
            parse_mode=ParseMode.HTML,
        )

        letter_opts = make_letter_only_options(len(raw_options))
        poll_kwargs = dict(
            chat_id=chat_id, question=".", options=letter_opts,
            type="quiz", correct_option_id=correct_index, is_anonymous=True,
        )
        if explanation:
            poll_kwargs["explanation"] = explanation[:TELEGRAM_EX_LIMIT]
        await _send_quiz_poll(context, poll_kwargs, image_path)

# ═══════════════════════════════════════════════════════════════
# PROGRESS MESSAGE BUILDER
# ═══════════════════════════════════════════════════════════════
def build_progress_text(items: list, latest_label: str = "") -> str:
    count   = len(items)
    bar_len = 4   # smaller block = the bar fills up faster (2 items = 50% full)

    if count == 0:
        filled = 0
    else:
        filled = count % bar_len or bar_len   # land on a full bar, not an empty one
    bar = "█" * filled + "░" * (bar_len - filled)

    type_counts = {"mcq": 0, "written": 0, "image": 0}
    for it in items:
        t = it.get("type", "mcq")
        if t in type_counts:
            type_counts[t] += 1

    breakdown = []
    if type_counts["mcq"]:
        breakdown.append(f"❓ {type_counts['mcq']} MCQ")
    if type_counts["written"]:
        breakdown.append(f"📝 {type_counts['written']} Written")
    if type_counts["image"]:
        breakdown.append(f"🖼 {type_counts['image']} Image")

    text = (
        f"📄 <b>PDF Collection Mode</b>\n"
        f"<code>{bar}</code>\n"
        f"Collected: <b>{count}</b> item{'s' if count != 1 else ''}"
    )
    if breakdown:
        text += f"\n{' · '.join(breakdown)}"
    return text

async def update_progress(context, user_id: int, chat_id: int, latest_label: str = ""):
    """Edit the existing progress message, or send a new one and store its id."""
    items    = PDF_BUFFER.get(user_id, [])
    text     = build_progress_text(items, latest_label)
    keyboard = export_keyboard()
    msg_id   = PROGRESS_MSG_ID.get(user_id)

    if msg_id:
        try:
            await context.bot.edit_message_text(
                chat_id=chat_id,
                message_id=msg_id,
                text=text,
                parse_mode=ParseMode.HTML,
                reply_markup=keyboard,
            )
            return
        except Exception:
            pass  # message too old / deleted — fall through to send new

    sent = await context.bot.send_message(
        chat_id=chat_id,
        text=text,
        parse_mode=ParseMode.HTML,
        reply_markup=keyboard,
    )
    PROGRESS_MSG_ID[user_id] = sent.message_id

# ═══════════════════════════════════════════════════════════════
# KEYBOARD HELPERS
# ═══════════════════════════════════════════════════════════════
def export_keyboard():
    row = [InlineKeyboardButton("📄 Export as PDF", callback_data="gen_pdf")]
    if DOCX_AVAILABLE:
        row.append(InlineKeyboardButton("📝 Export as DOCX", callback_data="gen_docx"))
    return InlineKeyboardMarkup([
        row,
        [InlineKeyboardButton("🗑 Clear & Cancel", callback_data="clear_pdf")],
    ])

def start_menu_keyboard():
    return InlineKeyboardMarkup([
        [
            InlineKeyboardButton("🦦 How To Use", callback_data="menu_how"),
            InlineKeyboardButton("Quizzes ⁉️",    callback_data="menu_quizzes"),
        ],
    ])

# ═══════════════════════════════════════════════════════════════
# MENU TEXT CONTENT
# ═══════════════════════════════════════════════════════════════
HOW_TO_USE_TEXT = (
    "📚 <b>How To Use — Quizician Bot</b>\n\n"
    "<b>1) Normal MCQ</b>\n"
    "<code>Question?\n"
    "a) Option A\n"
    "b) Option B z   ← mark correct with z\n"
    "c) Option C\n"
    "ex: Explanation here (optional)</code>\n\n"
    "<b>2) Single-line MCQ</b>\n"
    "<code>Question? a) A b) B z c) C</code>\n\n"
    "<b>3) Written / Flashcard</b>\n"
    "<code>Title\n"
    ".answer line 1\n"
    "answer line 2.</code>\n"
    "<i>Wrap the answer between dots.</i>\n\n"
    "<b>4) Forwarded Quiz Polls</b>\n"
    "Forward any Telegram quiz — the bot re-sends it with the correct answer preserved.\n\n"
    "<b>5) PDF / DOCX Mode</b>\n"
    "Use /pdf_start, collect items, then export.\n\n"
    "<b>6) 🤖 AI Extraction</b>\n"
    "Send a screenshot or PDF of MCQs with no caption — Quizzy reads it with AI, "
    "pulls out every question (even several at once), and shows you a preview "
    "to approve before anything's added. Always double-check the answers it picks!\n\n"
    "😴 /sleep — mute the bot until /start"
)

# ═══════════════════════════════════════════════════════════════
# PDF BUILDER
# ═══════════════════════════════════════════════════════════════
def build_pdf(items: list, doc_title: str = "questions") -> BytesIO:
    buffer = BytesIO()

    LEFT_COLOR  = colors.HexColor("#00BCD4")
    RIGHT_COLOR = colors.HexColor("#7B1FA2")

    def draw_header(canvas, doc):
        canvas.saveState()
        canvas.setFont(FONT_NAME_BOLD, 9)
        canvas.setFillColor(LEFT_COLOR)
        canvas.drawString(2 * cm, A4[1] - 1.4 * cm, "MDM44 | Notes & Files")
        canvas.setFillColor(RIGHT_COLOR)
        canvas.drawRightString(A4[0] - 2 * cm, A4[1] - 1.4 * cm, "Made by The Quizician")
        canvas.setStrokeColor(colors.HexColor("#CFD8DC"))
        canvas.setLineWidth(0.5)
        canvas.line(2 * cm, A4[1] - 1.65 * cm, A4[0] - 2 * cm, A4[1] - 1.65 * cm)
        canvas.restoreState()

    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm,
        topMargin=2.5*cm, bottomMargin=2*cm,
    )

    Q_STYLE = ParagraphStyle(
        "QStyle", fontName=FONT_NAME_BOLD, fontSize=12, leading=16,
        textColor=colors.HexColor("#1A1A2E"), spaceAfter=6, spaceBefore=14,
    )
    OPT_STYLE = ParagraphStyle(
        "OptStyle", fontName=FONT_NAME, fontSize=11, leading=15,
        textColor=colors.HexColor("#1A1A2E"), leftIndent=14, spaceAfter=3,
    )
    OPT_CORRECT = ParagraphStyle(
        "OptCorrect", fontName=FONT_NAME_BOLD, fontSize=11, leading=15,
        textColor=colors.HexColor("#1B5E20"), leftIndent=14, spaceAfter=3,
    )
    WRITTEN_TITLE = ParagraphStyle(
        "WTitle", fontName=FONT_NAME_BOLD, fontSize=12, leading=16,
        textColor=colors.HexColor("#1A1A2E"), spaceAfter=4, spaceBefore=14,
    )
    WRITTEN_BODY = ParagraphStyle(
        "WBody", fontName=FONT_NAME, fontSize=11, leading=15,
        textColor=colors.HexColor("#37474F"), leftIndent=14, spaceAfter=6,
    )
    NUM_STYLE = ParagraphStyle(
        "NumStyle", fontName=FONT_NAME_BOLD, fontSize=9,
        textColor=colors.HexColor("#90A4AE"), spaceAfter=2,
    )
    IMG_CAPTION = ParagraphStyle(
        "ImgCaption", fontName=FONT_NAME, fontSize=9, leading=12,
        textColor=colors.HexColor("#78909C"), spaceAfter=6, spaceBefore=4,
    )

    HR_COLOR = colors.HexColor("#CFD8DC")
    story    = []

    for idx, item in enumerate(items, 1):
        q_num_label = f"~Q{idx}" if item.get("type") == "mcq" and item.get("correct") is None else f"Q{idx}"
        story.append(Paragraph(q_num_label, NUM_STYLE))

        if item["type"] == "mcq":
            story.append(Paragraph(item["q"], Q_STYLE))
            if item.get("image"):
                try:
                    img = RLImage(item["image"])
                    if img.imageWidth > PDF_MAX_IMG_WIDTH:
                        scale          = PDF_MAX_IMG_WIDTH / img.imageWidth
                        img.drawWidth  = PDF_MAX_IMG_WIDTH
                        img.drawHeight = img.imageHeight * scale
                    story.append(Spacer(1, 6))
                    story.append(img)
                    story.append(Spacer(1, 6))
                except Exception as e:
                    story.append(Paragraph(f"[Image error: {e}]", WRITTEN_BODY))
            for i, opt in enumerate(item["options"]):
                if i == item["correct"]:
                    story.append(Paragraph(f"✓  {opt}", OPT_CORRECT))
                else:
                    story.append(Paragraph(f"     {opt}", OPT_STYLE))

        elif item["type"] == "written":
            story.append(Paragraph(item["title"], WRITTEN_TITLE))
            for line in item["content"].split("\n"):
                line = line.strip()
                if line:
                    story.append(Paragraph(f"• {line}", WRITTEN_BODY))

        elif item["type"] == "image":
            img_path = item["path"]
            try:
                img = RLImage(img_path)
                if img.imageWidth > PDF_MAX_IMG_WIDTH:
                    scale          = PDF_MAX_IMG_WIDTH / img.imageWidth
                    img.drawWidth  = PDF_MAX_IMG_WIDTH
                    img.drawHeight = img.imageHeight * scale
                story.append(Spacer(1, 8))
                story.append(img)
                if item.get("caption"):
                    story.append(Paragraph(f"📷 {item['caption']}", IMG_CAPTION))
                story.append(Spacer(1, 4))
            except Exception as e:
                story.append(Paragraph(f"[Image error: {e}]", WRITTEN_BODY))

        if idx < len(items):
            story.append(Spacer(1, 6))
            story.append(HRFlowable(width="100%", thickness=0.5, color=HR_COLOR, spaceAfter=4))

    doc.build(story, onFirstPage=draw_header, onLaterPages=draw_header)
    buffer.seek(0)
    return buffer

# ═══════════════════════════════════════════════════════════════
# DOCX BUILDER  — pure Python, no Node.js
# ═══════════════════════════════════════════════════════════════
def _hex_to_rgb(hex_color: str):
    """Convert 'RRGGBB' string to RGBColor."""
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

def _add_paragraph(doc, text: str, bold=False, size_pt=11,
                   color_hex="1A1A2E", indent_cm=0,
                   space_before=0, space_after=6,
                   align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    p   = doc.add_paragraph()
    p.alignment = align
    pf  = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if indent_cm:
        pf.left_indent = Cm(indent_cm)
    run = p.add_run(text)
    run.bold        = bold
    run.font.size   = Pt(size_pt)
    run.font.color.rgb = _hex_to_rgb(color_hex)
    return p

def _add_horizontal_rule(doc):
    """Add a thin bottom border to simulate a horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CFD8DC")
    pBdr.append(bottom)
    pPr.append(pBdr)

def _set_header_border(para):
    """Add bottom border to the header paragraph."""
    pPr   = para._p.get_or_add_pPr()
    pBdr  = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "CFD8DC")
    pBdr.append(bottom)
    pPr.append(pBdr)

def build_docx(items: list, doc_title: str = "questions") -> BytesIO:
    doc = DocxDocument()

    # ── Page margins ──────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # ── Header ────────────────────────────────────────────────
    header_para = doc.add_paragraph()
    header_para.paragraph_format.space_after = Pt(8)
    _set_header_border(header_para)

    r1 = header_para.add_run("MDM44 | Notes & Files")
    r1.bold            = True
    r1.font.size       = Pt(9)
    r1.font.color.rgb  = _hex_to_rgb("00BCD4")

    header_para.add_run("   ·   ")

    r2 = header_para.add_run("Made by The Quizician")
    r2.bold            = True
    r2.font.size       = Pt(9)
    r2.font.color.rgb  = _hex_to_rgb("7B1FA2")

    # ── Items ─────────────────────────────────────────────────
    for idx, item in enumerate(items, 1):

        # Q-number label
        q_num_label = f"~Q{idx}" if item.get("type") == "mcq" and item.get("correct") is None else f"Q{idx}"
        _add_paragraph(doc, q_num_label, bold=True, size_pt=8,
                       color_hex="90A4AE", space_before=10, space_after=2)

        if item["type"] == "mcq":
            _add_paragraph(doc, item["q"], bold=True, size_pt=12,
                           color_hex="1A1A2E", space_before=0, space_after=4)
            if item.get("image") and os.path.exists(item["image"]):
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after  = Pt(6)
                    run = p.add_run()
                    run.add_picture(item["image"], width=Inches(5.5))
                except Exception as e:
                    _add_paragraph(doc, f"[Image error: {e}]",
                                   size_pt=10, color_hex="B71C1C")
            for i, opt in enumerate(item["options"]):
                correct = (i == item["correct"])
                _add_paragraph(
                    doc,
                    ("✓  " if correct else "     ") + opt,
                    bold=correct, size_pt=11,
                    color_hex="1B5E20" if correct else "1A1A2E",
                    indent_cm=0.7, space_after=3,
                )

        elif item["type"] == "written":
            _add_paragraph(doc, item["title"], bold=True, size_pt=12,
                           color_hex="1A1A2E", space_before=0, space_after=4)
            for line in item["content"].split("\n"):
                line = line.strip()
                if line:
                    _add_paragraph(doc, f"• {line}", bold=False, size_pt=11,
                                   color_hex="37474F", indent_cm=0.7, space_after=3)

        elif item["type"] == "image":
            img_path = item.get("path", "")
            if img_path and os.path.exists(img_path):
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after  = Pt(4)
                    run = p.add_run()
                    run.add_picture(img_path, width=Inches(5.5))
                    if item.get("caption"):
                        cap = _add_paragraph(
                            doc, f"📷 {item['caption']}",
                            bold=False, size_pt=9, color_hex="78909C",
                            space_after=4,
                        )
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    _add_paragraph(doc, f"[Image error: {e}]",
                                   size_pt=10, color_hex="B71C1C")
            else:
                _add_paragraph(doc, "[Image file not found]",
                               size_pt=10, color_hex="B71C1C")

        # Divider between items
        if idx < len(items):
            _add_horizontal_rule(doc)

    # ── Save to BytesIO ───────────────────────────────────────
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ═══════════════════════════════════════════════════════════════
# REACTIONS
# ═══════════════════════════════════════════════════════════════
async def react_random(update: Update, context: ContextTypes.DEFAULT_TYPE):
    try:
        roll  = random.randint(1, 20)
        emoji = "🫡" if roll <= 15 else "❤️" if roll <= 19 else "🏆"
        await context.bot.set_message_reaction(
            chat_id=update.effective_chat.id,
            message_id=update.message.message_id,
            reaction=[ReactionTypeEmoji(emoji)],
            is_big=False,
        )
    except Exception:
        pass

# ═══════════════════════════════════════════════════════════════
# SLEEP / WAKE COMMANDS
# ═══════════════════════════════════════════════════════════════
async def sleep_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_chat.id
    SLEEPING.add(user_id)
    await update.message.reply_text(
        f"{quizzy_block(QUIZZY_SLEEPING_ART, 'قوزي نام، وأنا نايم معاه 😴')}\n\n"
        "نادينا بـ /start لما تحتاجنا تاني",
        parse_mode=ParseMode.HTML,
    )

# ═══════════════════════════════════════════════════════════════
# PASSIVE ANSWER BACKFILL
# Telegram pushes a fresh Update.poll (with correct_option_ids filled in)
# to any bot that has previously seen a poll, once that poll is stopped —
# even for polls the bot didn't create. If the original quiz's creator
# later ends it, we quietly backfill the answer with no user action needed.
# ═══════════════════════════════════════════════════════════════
async def poll_update_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    poll = update.poll

    # ── Quiz-channel poll tracking: mark it closed once stopped ──
    if poll is not None and poll.id in QUIZ_POLL_STATUS and poll.is_closed:
        entry = QUIZ_POLL_STATUS[poll.id]
        if not entry["closed"]:
            entry["closed"] = True
            if poll.correct_option_ids and entry.get("correct_option_id") is None:
                entry["correct_option_id"] = poll.correct_option_ids[0]
            # Capture full poll content too, not just the correct answer —
            # lecture delivery needs to rebuild these as our own polls (see
            # _deliver_next_lecture_question) so it can get real poll_answer
            # events, since a straight copy of a channel poll stays
            # anonymous forever and Telegram never sends those for it.
            entry["question"]    = poll.question
            entry["options"]     = [o.text for o in poll.options]
            entry["explanation"] = poll.explanation
            save_quiz_poll_status()
            try:
                await context.bot.set_message_reaction(
                    chat_id=QUIZ_CHANNEL_ID, message_id=entry["message_id"],
                    reaction=[ReactionTypeEmoji("✅")], is_big=False,
                )
            except Exception:
                pass

    if poll is None or not poll.correct_option_ids:
        return

    watch = POLL_WATCH.pop(poll.id, None)
    if not watch:
        return
    user_id, item_index = watch

    items = PDF_BUFFER.get(user_id)
    if not items or item_index >= len(items) or items[item_index]["correct"] is not None:
        return  # buffer changed, or already resolved manually — skip

    item = items[item_index]
    correct_id = poll.correct_option_ids[0]
    item["correct"] = correct_id

    queue = CLARIFY_QUEUE.get(user_id, [])
    if item_index in queue:
        queue.remove(item_index)

    try:
        await context.bot.send_message(
            chat_id=user_id,
            text=(
                f"✅ الكويز الأصلي لسؤال Q{item_index + 1} اتقفل وتليجرام بعت الإجابة الصح تلقائي: "
                f"{item['options'][correct_id]}"
            ),
        )
    except Exception:
        pass

# ── XP for lecture quiz answers ───────────────────────────────
XP_LECTURE_QUESTION       = 10   # flat, per question answered — right or wrong doesn't matter
XP_LECTURE_COMPLETE_BONUS = 25   # extra, on top of the above, for the lecture's last question

async def _deliver_next_lecture_question(context: ContextTypes.DEFAULT_TYPE, user_id: int, session: dict) -> bool:
    """Pops message ids off session['queue'] and sends them to the user one
    at a time as the bot's own non-anonymous quiz polls (built from content
    captured off the channel poll when it closed — see poll_update_handler),
    skipping (and cleaning up) any that were deleted from the channel since
    indexing, until one is delivered or the queue runs dry.

    Why not just copy_message the original? Because channel polls (and
    therefore any straight copy of one) are always anonymous, and Telegram
    never sends poll_answer for anonymous polls — there'd be no way to tell
    the question had been answered. Sending our own copy with
    is_anonymous=False sidesteps that entirely.

    For lecture questions closed before this content-capture existed, the
    QUIZ_POLL_STATUS entry won't have "question"/"options" yet — those get
    backfilled here via a one-time throwaway forward (read its poll
    content, delete it, never shown to the user) the first time they're
    delivered.

    Sets session['current_poll_id']. Returns whether a question went out."""
    entry = QUIZ_INDEX.get(session["lecture_key"])

    def _drop_dead(mid: int):
        if entry and mid in entry.get("ids", []):
            entry["ids"].remove(mid)
            if not entry["ids"]:
                QUIZ_INDEX.pop(session["lecture_key"], None)  # whole lecture was deleted
        for pid in [pid for pid, v in QUIZ_POLL_STATUS.items() if v["message_id"] == mid]:
            QUIZ_POLL_STATUS.pop(pid, None)
        save_quiz_index()
        save_quiz_poll_status()

    while session["queue"]:
        mid = session["queue"].pop(0)
        status = next((v for v in QUIZ_POLL_STATUS.values() if v["message_id"] == mid), None)

        question    = status.get("question")    if status else None
        options     = status.get("options")      if status else None
        correct_id  = status.get("correct_option_id") if status else None
        explanation = status.get("explanation")  if status else None

        if not (question and options and correct_id is not None):
            # Legacy entry (closed before content-capture existed) — grab
            # the content via a throwaway forward, then delete it; the
            # forward itself is never what gets answered. Must be
            # forward_message here, not copy_message: copyMessage's API
            # response is just a bare message_id with no poll content at
            # all, so there'd be nothing here to read.
            try:
                probe = await context.bot.forward_message(chat_id=user_id, from_chat_id=QUIZ_CHANNEL_ID, message_id=mid)
            except Exception as e:
                print(f"Quiz question {mid} in lecture '{session['lecture_key']}' unreachable (likely deleted): {e}")
                _drop_dead(mid)
                continue
            if probe.poll and probe.poll.correct_option_ids:
                question    = probe.poll.question
                options     = [o.text for o in probe.poll.options]
                correct_id  = probe.poll.correct_option_ids[0]
                explanation = probe.poll.explanation
                if status is not None:
                    status.update(question=question, options=options,
                                   correct_option_id=correct_id, explanation=explanation)
                    save_quiz_poll_status()
            try:
                await context.bot.delete_message(chat_id=user_id, message_id=probe.message_id)
            except Exception:
                pass
            if not (question and options and correct_id is not None):
                continue  # still couldn't recover real quiz content — skip it

        try:
            msg = await context.bot.send_poll(
                chat_id=user_id, question=question, options=options,
                type="quiz", correct_option_id=correct_id, is_anonymous=False,
                explanation=(explanation or None),
            )
        except Exception as e:
            print(f"Couldn't send lecture question {mid}: {e}")
            continue

        session["current_poll_id"] = msg.poll.id
        return True

    session["current_poll_id"] = None
    return False

async def handle_poll_answer(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Fires when a user answers a poll the bot sent (our lecture-delivery
    polls are always is_anonymous=False specifically so this reliably
    fires). If it's the question we're currently waiting on for that user's
    active lecture session, hand off to _advance_lecture_session."""
    answer  = update.poll_answer
    poll_id = answer.poll_id
    user_id = answer.user.id

    session = LECTURE_SESSIONS.get(user_id)
    if not session or session.get("current_poll_id") != poll_id:
        return   # not the question we're tracking for this user right now

    await _advance_lecture_session(context, user_id, session)


async def _advance_lecture_session(context: ContextTypes.DEFAULT_TYPE, user_id: int, session: dict):
    """Called once handle_poll_answer confirms the user answered their
    current lecture question. Awards flat XP — right or wrong makes no
    difference — and immediately forwards the next question, with a bonus
    tacked on once the lecture runs out."""
    session["answered"] += 1
    sent_next = await _deliver_next_lecture_question(context, user_id, session)
    is_last   = not sent_next   # queue ran dry (or every remaining id was dead) — lecture's done

    xp_delta   = XP_LECTURE_QUESTION + (XP_LECTURE_COMPLETE_BONUS if is_last else 0)
    events     = _record_activity(user_id)
    user_entry = _get_entry(user_id)
    new_level  = _award_xp(user_entry, xp_delta)
    if new_level:
        events["level_up"] = new_level
    save_analytics()
    await _announce_events(context, user_id, events)

    try:
        if is_last:
            text = (
                f"🎉 خلصت المحاضرة! <b>+{XP_LECTURE_QUESTION} XP</b>، "
                f"وبونص إتمام المحاضرة <b>+{XP_LECTURE_COMPLETE_BONUS} XP</b> "
                f"→ إجمالي <b>+{xp_delta} XP</b>"
            )
        else:
            text = f"<b>+{xp_delta} XP</b> ⭐"
        await context.bot.send_message(chat_id=user_id, text=text, parse_mode=ParseMode.HTML)
    except Exception:
        pass

    await backup_analytics_to_channel(context)

    if is_last:
        LECTURE_SESSIONS.pop(user_id, None)


# ═══════════════════════════════════════════════════════════════
# FORWARDED POLL HANDLER
# ═══════════════════════════════════════════════════════════════
async def _ask_next_clarification(context, user_id: int, chat_id: int):
    """Pop-free peek at the front of the clarify queue and ask about it with
    inline A/B/C… buttons. Skips (and drops) any stale entries whose buffer
    item no longer exists (e.g. buffer was cleared mid-queue)."""
    queue = CLARIFY_QUEUE.get(user_id)
    while queue:
        item_index = queue[0]
        items = PDF_BUFFER.get(user_id)
        if not items or item_index >= len(items) or items[item_index]["correct"] is not None:
            queue.pop(0)  # stale or already resolved — skip it
            continue

        item        = items[item_index]
        q_num       = item_index + 1
        first_words = " ".join(item["q"].split()[:5])
        options_txt = "\n".join(item["options"])  # already "A) ..." labeled

        buttons = [
            InlineKeyboardButton(string.ascii_uppercase[i], callback_data=f"clarify:{item_index}:{i}")
            for i in range(len(item["options"]))
        ]
        rows = [buttons[i:i + 6] for i in range(0, len(buttons), 6)]

        await context.bot.send_message(
            chat_id=chat_id,
            text=(
                f"❓ <b>Choose the correct answer</b>\n"
                f"for Q{q_num}: {first_words}…\n\n{options_txt}"
            ),
            parse_mode=ParseMode.HTML,
            reply_markup=InlineKeyboardMarkup(rows),
        )
        return
    # queue exhausted — nothing left to ask
    CLARIFY_QUEUE.pop(user_id, None)

# ═══════════════════════════════════════════════════════════════
# QUESTION REVIEW / EDIT — after a question lands in the PDF buffer
# (correct answer already known at this point, whether that came in
# automatically or via the clarify buttons above), show it back to the
# admin with buttons to tweak the question text, any option's text, or
# even flip which option is correct — before moving on.
# ═══════════════════════════════════════════════════════════════
def _review_text(item: dict) -> str:
    if item["type"] == "written":
        return f"📝 <b>{html.escape(item['title'])}</b>\n{html.escape(item['content'])}"
    lines = [f"❓ {html.escape(item['q'])}"]
    for i, opt in enumerate(item["options"]):
        mark = "  ✅" if i == item.get("correct") else ""
        lines.append(html.escape(opt) + mark)
    return "\n".join(lines)

def _review_buttons(item_index: int, item: dict) -> InlineKeyboardMarkup:
    if item["type"] == "written":
        rows = [
            [InlineKeyboardButton("✏️ عدّل العنوان", callback_data=f"revedit:{item_index}:title")],
            [InlineKeyboardButton("✏️ عدّل المحتوى", callback_data=f"revedit:{item_index}:content")],
            [InlineKeyboardButton("✅ تمام، مفيش تعديل", callback_data=f"revedit:{item_index}:done")],
        ]
        return InlineKeyboardMarkup(rows)

    opt_buttons = [
        InlineKeyboardButton(f"✏️ {string.ascii_uppercase[i]}", callback_data=f"revedit:{item_index}:opt:{i}")
        for i in range(len(item["options"]))
    ]
    rows = [opt_buttons[i:i + 6] for i in range(0, len(opt_buttons), 6)]
    rows.append([InlineKeyboardButton("✏️ عدّل نص السؤال", callback_data=f"revedit:{item_index}:q")])
    if item.get("correct") is not None:
        rows.append([InlineKeyboardButton("🔁 غيّر الإجابة الصح", callback_data=f"revedit:{item_index}:correct")])
    rows.append([InlineKeyboardButton("✅ تمام، مفيش تعديل", callback_data=f"revedit:{item_index}:done")])
    return InlineKeyboardMarkup(rows)

def _edit_button_markup(item_index: int) -> InlineKeyboardMarkup:
    """Single '✏️ تعديل' button — shown on confirmation, pressed only if needed."""
    return InlineKeyboardMarkup([[
        InlineKeyboardButton("✏️ تعديل", callback_data=f"revedit:{item_index}:open")
    ]])

async def handle_poll(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not update.message or not update.message.poll:
        return

    user_id = update.effective_chat.id
    if user_id in SLEEPING:
        return

    poll = update.message.poll
    # Strip any existing A) B) C) prefixes from options to avoid double-labeling
    question      = poll.question
    raw_options   = [strip_leading_letter_prefix(opt.text) for opt in poll.options]
    # Telegram only reveals correct_option_ids if the quiz is closed, or was
    # sent by our own bot / directly to it — an open quiz forwarded from
    # someone else comes back empty. We must NOT guess in that case.
    correct_index = poll.correct_option_ids[0] if poll.correct_option_ids else None
    explanation   = poll.explanation or None

    # An image sent (with no caption / unparseable caption) just before this
    # forward is paired with it, in either mode.
    pending_img = PENDING_IMAGE.pop(user_id, None)

    # ── PDF mode: save poll (+ any paired image) to buffer ──────
    if user_id in PDF_BUFFER:
        labeled_options = [
            f"{string.ascii_uppercase[i]}) {opt}" for i, opt in enumerate(raw_options)
        ]
        item = {
            "type": "mcq", "q": question,
            "options": labeled_options, "correct": correct_index,  # None = unknown
            "poll_id": poll.id,
        }
        if pending_img:
            item["image"] = pending_img

        PDF_BUFFER[user_id].append(item)
        item_index = len(PDF_BUFFER[user_id]) - 1

        label = ("🖼 " if pending_img else "") + ("~" if correct_index is None else "") \
                + question[:50] + ("…" if len(question) > 50 else "")
        await update_progress(context, user_id, update.effective_chat.id, latest_label=label)

        if correct_index is None:
            # Telegram hid the answer (quiz still open, not ours) — queue it
            # for a quick button tap instead of silently guessing. The
            # review/edit prompt fires once that tap resolves it (see the
            # "clarify:" branch in button_handler).
            POLL_WATCH[poll.id] = (user_id, item_index)
            queue = CLARIFY_QUEUE.setdefault(user_id, [])
            queue.append(item_index)
            if len(queue) == 1:  # nothing else currently being asked
                await _ask_next_clarification(context, user_id, update.effective_chat.id)
        else:
            # Correct answer already known — show confirmation with edit button.
            short = question[:50] + ("…" if len(question) > 50 else "")
            await context.bot.send_message(
                chat_id=update.effective_chat.id,
                text=f"✅ اتسجل: {html.escape(short)}",
                parse_mode=ParseMode.HTML,
                reply_markup=_edit_button_markup(item_index),
            )
        return

    # ── Normal mode: show the poll question + choices with an edit button ──
    lines = [f"❓ <b>{html.escape(question)}</b>"]
    for opt in raw_options:
        lines.append(html.escape(opt))
    full_text = "\n".join(lines)

    # Build a temporary "normal mode" item so the edit flow works the same way
    normal_item = {
        "type": "mcq", "q": question,
        "options": [f"{string.ascii_uppercase[i]}) {o}" if not o.startswith(tuple(string.ascii_uppercase)) else o
                    for i, o in enumerate(raw_options)],
        "correct": None,
    }
    nm_buf = PDF_BUFFER.setdefault(user_id, [])
    nm_buf.append(normal_item)
    nm_index = len(nm_buf) - 1

    if pending_img:
        with open(pending_img, "rb") as f:
            cap = full_text if len(full_text) <= 1024 else None
            sent = await context.bot.send_photo(
                chat_id=user_id, photo=f, caption=cap,
                parse_mode=ParseMode.HTML if cap else None,
                reply_markup=_edit_button_markup(nm_index) if cap else None,
            )
            if not cap:
                await context.bot.send_message(
                    chat_id=user_id, text=full_text,
                    parse_mode=ParseMode.HTML,
                    reply_markup=_edit_button_markup(nm_index),
                )
    else:
        await context.bot.send_message(
            chat_id=user_id, text=full_text,
            parse_mode=ParseMode.HTML,
            reply_markup=_edit_button_markup(nm_index),
        )

# ═══════════════════════════════════════════════════════════════
# IMAGE HANDLER (PDF mode only)
# ═══════════════════════════════════════════════════════════════
async def handle_image(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not update.message:
        return

    user_id  = update.effective_chat.id
    real_uid = update.effective_user.id if update.effective_user else user_id
    if user_id in SLEEPING:
        return

    photo = update.message.photo[-1] if update.message.photo else None
    if not photo:
        return

    in_pdf_mode = user_id in PDF_BUFFER
    caption     = (update.message.caption or "").strip()

    # ── Download the image (works in both PDF and normal mode now) ──
    img_dir = os.path.join(IMG_BASE_DIR, str(user_id))
    os.makedirs(img_dir, exist_ok=True)
    img_path = os.path.join(img_dir, f"img_{photo.file_unique_id}.jpg")
    tg_file  = await context.bot.get_file(photo.file_id)
    await tg_file.download_to_drive(img_path)

    # ── Case 1: caption already IS a complete quiz question ─────────
    # Parse and build the question immediately — no need to ask again.
    parsed = parse_mcq_block(caption) if caption else None
    if parsed:
        question, raw_options, correct_index, explanation = parsed
        if in_pdf_mode:
            labeled_options = [
                f"{string.ascii_uppercase[i]}) {opt}" for i, opt in enumerate(raw_options)
            ]
            PDF_BUFFER[user_id].append({
                "type": "mcq", "q": question,
                "options": labeled_options, "correct": correct_index,
                "image": img_path,
            })
            await update_progress(
                context, user_id, update.effective_chat.id,
                latest_label=f"🖼 {question[:50]}" + ("…" if len(question) > 50 else ""),
            )
        else:
            await deliver_quiz(
                context, user_id, question, raw_options, correct_index,
                explanation=explanation, image_path=img_path,
            )
            events = _record_activity(real_uid, questions_delta=1)
            await react_random(update, context)
            await _announce_events(context, user_id, events)
            await backup_analytics_to_channel(context)
        return

    # ── Case 2 (PDF mode only): non-empty caption that ISN'T a full
    # question — keep the old behaviour of saving it as a standalone
    # image item (e.g. comparison charts / tables with a plain caption).
    if in_pdf_mode and caption:
        PDF_BUFFER[user_id].append({
            "type": "image", "path": img_path, "caption": caption,
        })
        await update_progress(
            context, user_id, update.effective_chat.id,
            latest_label=f"Image — {caption}",
        )
        return

    # ── Case 3: no caption — park the image and ask for the question
    _clear_pending_image(user_id)
    PENDING_IMAGE[user_id] = img_path
    await update.message.reply_text(
        "🖼 <b>استلمت الصورة!</b>\n"
        "دلوقتي ابعت السؤال والاختيارات (بنفس صيغة الأسئلة المعتادة) "
        "وهيتضاف الصورة تلقائي للسؤال ده.",
        parse_mode=ParseMode.HTML,
    )

async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """PDFs sent in a private DM: captioned = treated as a manual question
    (caption parsed as the MCQ text — no attachment, since Telegram polls
    can only carry a photo, not a PDF). Uncaptioned PDFs are silently ignored."""
    if not update.message or not update.message.document:
        return
    user_id  = update.effective_chat.id
    real_uid = update.effective_user.id if update.effective_user else user_id
    if user_id in SLEEPING:
        return
    doc = update.message.document
    if doc.mime_type != "application/pdf":
        return

    caption = (update.message.caption or "").strip()
    if caption:
        in_pdf_mode = user_id in PDF_BUFFER
        parsed = parse_mcq_block(caption)
        if not parsed:
            await update.message.reply_text(
                "⚠️ الكابشن مش صيغة سؤال كاملة (لازم سؤال + اختيارات + إجابة صح متعلّم عليها بـ z).\n"
                "ابعت الملف تاني من غير كابشن لو عايز تستخدم استخراج الـ AI بدل كده."
            )
            return
        question, raw_options, correct_index, explanation = parsed
        if in_pdf_mode:
            labeled_options = [f"{string.ascii_uppercase[i]}) {opt}" for i, opt in enumerate(raw_options)]
            PDF_BUFFER[user_id].append({"type": "mcq", "q": question, "options": labeled_options, "correct": correct_index})
            await update_progress(
                context, user_id, update.effective_chat.id,
                latest_label=f"📄 {question[:50]}" + ("…" if len(question) > 50 else ""),
            )
        else:
            await deliver_quiz(context, user_id, question, raw_options, correct_index, explanation=explanation)
            events = _record_activity(real_uid, questions_delta=1)
            await react_random(update, context)
            await _announce_events(context, user_id, events)
            await backup_analytics_to_channel(context)
        return


# ═══════════════════════════════════════════════════════════════
# STORAGE GROUP — AUTO-INDEXING
# ═══════════════════════════════════════════════════════════════
def _index_item(caption: str, message_ids: list):
    password = caption.strip().split(maxsplit=1)[0].lower()
    STORAGE_INDEX.setdefault(password, []).append(sorted(message_ids))
    save_storage_index()
    return password

async def _finalize_album(context: ContextTypes.DEFAULT_TYPE, media_group_id: str):
    # Wait for the album's parts to stop arriving before filing it as one item.
    await asyncio.sleep(1.5)
    buf = ALBUM_BUFFER.pop(media_group_id, None)
    if not buf:
        return
    caption = buf["caption"]
    if not caption:
        await context.bot.send_message(
            STORAGE_GROUP_ID,
            "⚠️ ألبوم اتبعت من غير كابشن (كلمة سر) — اتجاهله ومحدش هيقدر يفتحه.",
        )
        return
    password = _index_item(caption, buf["ids"])
    await backup_storage_to_channel(context)
    await context.bot.send_message(
        STORAGE_GROUP_ID,
        f"✅ اتخزن ألبوم من {len(buf['ids'])} ملف تحت الكلمة: <code>{password}</code>\n"
        f"🐾 <i>{random.choice(QUIZZY_SUCCESS_LINES)}</i>",
        parse_mode=ParseMode.HTML,
    )

async def handle_storage_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Indexes media posted in STORAGE_GROUP_ID. First caption word = password."""
    msg = update.message
    if not msg:
        return
    caption = (msg.caption or "").strip()

    if msg.media_group_id:
        buf = ALBUM_BUFFER.setdefault(msg.media_group_id, {"ids": [], "caption": None})
        buf["ids"].append(msg.message_id)
        if caption:
            buf["caption"] = caption  # usually only one part of the album carries it
        existing_task = buf.get("task")
        if existing_task:
            existing_task.cancel()
        buf["task"] = asyncio.create_task(_finalize_album(context, msg.media_group_id))
        return

    if not caption:
        await msg.reply_text("⚠️ الملف ده اتبعت من غير كابشن — محتاج كلمة سر في الكابشن عشان يتخزن.")
        return

    password = _index_item(caption, [msg.message_id])
    await backup_storage_to_channel(context)
    await msg.reply_text(
        f"✅ اتخزن تحت الكلمة: <code>{password}</code>\n"
        f"🐾 <i>{random.choice(QUIZZY_SUCCESS_LINES)}</i>",
        parse_mode=ParseMode.HTML,
    )

async def storage_id_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Utility: run inside the storage group to get its chat ID for STORAGE_GROUP_ID."""
    await update.message.reply_text(
        f"🆔 Chat ID: <code>{update.effective_chat.id}</code>", parse_mode=ParseMode.HTML
    )

async def backup_now_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Admin: force-create/refresh both pinned backups right now, instead of
    waiting for the next real change."""
    if not is_admin(update):
        await update.message.reply_text(MSG_ADMIN_ONLY)
        return
    await backup_storage_to_channel(context)
    await backup_quiz_to_channel(context)
    await update.message.reply_text(
        "✅ اتعمل باك أب دلوقتي.\n"
        f"📌 Storage group: {'تم' if STORAGE_BACKUP_STATE.get('backup_msg_id') else 'مش متظبط STORAGE_GROUP_ID'}\n"
        f"📌 Quiz channel: {'تم' if QUIZ_BACKUP_STATE.get('backup_msg_id') else 'مش متظبط QUIZ_CHANNEL_ID'}"
    )

# ═══════════════════════════════════════════════════════════════
# QUIZ CHANNEL — AUTO-INDEXING
# ═══════════════════════════════════════════════════════════════
async def handle_quiz_channel_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Indexes lectures + quiz polls posted in QUIZ_CHANNEL_ID."""
    msg = update.channel_post or update.message
    if not msg:
        return

    # ── A quiz poll — file it under the currently-open lecture ──
    if msg.poll:
        current = QUIZ_STATE.get("current_lecture")
        if not current or current not in QUIZ_INDEX:
            await context.bot.send_message(
                QUIZ_CHANNEL_ID,
                "⚠️ محتاج تبعت اسم المحاضرة الأول (أي رسالة نصية) قبل ما تبعت أسئلة."
            )
            return
        QUIZ_INDEX[current]["ids"].append(msg.message_id)
        save_quiz_index()
        # Track this poll so we know once it's stopped (only then is the
        # correct answer known — needed before it can be delivered as a
        # lecture question). question/options are captured right away since
        # those aren't access-restricted like correct_option_id is; that
        # lets lecture delivery build its own poll (see
        # _deliver_next_lecture_question) without depending on a copy of
        # the original, which would stay anonymous forever.
        QUIZ_POLL_STATUS[msg.poll.id] = {
            "lecture":           current,
            "message_id":        msg.message_id,
            "closed":            msg.poll.is_closed,
            "correct_option_id": msg.poll.correct_option_ids[0] if msg.poll.correct_option_ids else None,
            "question":          msg.poll.question,
            "options":           [o.text for o in msg.poll.options],
            "explanation":       msg.poll.explanation,
        }
        save_quiz_poll_status()
        # NOTE: the channel backup document + the "still open" reaction are
        # both deliberately deferred to -END (below) instead of happening
        # here per-question — doing them per-question was sending/pinning
        # a fresh backup document for every single poll.
        return

    # ── Plain text: either "-END" or a new/resumed lecture name ─
    if not msg.text:
        return
    text = msg.text.strip()

    if text.upper() == "-END":
        current = QUIZ_STATE.get("current_lecture")
        if not current or current not in QUIZ_INDEX:
            await context.bot.send_message(QUIZ_CHANNEL_ID, "⚠️ مفيش محاضرة مفتوحة دلوقتي.")
            return
        QUIZ_INDEX[current]["closed"] = True
        save_quiz_index()
        QUIZ_STATE["current_lecture"] = None
        save_quiz_state()

        # Batched now, once, instead of one reaction call per question:
        # mark every still-open (forgot to Stop Poll) question in this
        # lecture with 😢.
        open_message_ids = [
            p["message_id"] for p in QUIZ_POLL_STATUS.values()
            if p["lecture"] == current and not p["closed"]
        ]
        for mid in open_message_ids:
            try:
                await context.bot.set_message_reaction(
                    chat_id=QUIZ_CHANNEL_ID, message_id=mid,
                    reaction=[ReactionTypeEmoji("😢")], is_big=False,
                )
            except Exception:
                pass

        # Single backup for the whole lecture, once it's actually closed.
        await backup_quiz_to_channel(context)

        count = len(QUIZ_INDEX[current]["ids"])
        open_count = len(open_message_ids)
        note = (
            f"\n⚠️ {open_count} سؤال لسه مفتوح — لازم توقف التصويت عليه (Stop Poll) "
            f"قبل ما يبقى ممكن يتبعت للطلاب."
            if open_count else "\n✅ كل الأسئلة جاهزة للإرسال."
        )
        await context.bot.send_message(
            QUIZ_CHANNEL_ID,
            f"✅ اتقفلت محاضرة <b>{current}</b> — {count} سؤال.{note}",
            parse_mode=ParseMode.HTML,
        )
        return

    # New lecture name (or resuming one that already exists).
    # Format: "<Module> - <Subject> Lecture <number>: <name>"
    module, subject, lecture_number, name_or_error = parse_lecture_title(text)
    if module is None:
        await context.bot.send_message(QUIZ_CHANNEL_ID, name_or_error, parse_mode=ParseMode.HTML)
        return
    name = name_or_error

    entry = QUIZ_INDEX.setdefault(text, {
        "ids": [], "closed": False,
        "module": module, "subject": subject, "lecture_number": lecture_number, "name": name,
    })
    entry["closed"]         = False
    entry["module"]         = module
    entry["subject"]        = subject
    entry["lecture_number"] = lecture_number
    entry["name"]           = name
    save_quiz_index()
    QUIZ_STATE["current_lecture"] = text
    save_quiz_state()
    await backup_quiz_to_channel(context)
    await context.bot.send_message(
        QUIZ_CHANNEL_ID,
        f"🆕 <b>{module} - {subject} Lecture {lecture_number}: {name}</b>\n"
        f"ابعت الأسئلة (كويزات) دلوقتي، وابعت <code>-END</code> لما تخلص.\n"
        f"⚠️ لازم توقف كل سؤال (Stop Poll) قبل الـ -END عشان يبقى قابل للإرسال.",
        parse_mode=ParseMode.HTML,
    )

async def quiz_channel_id_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Utility: forward any message from the quiz channel here first, then
    run this command in the same DM — it reads the forward's source chat ID."""
    fwd = update.message.forward_from_chat if update.message else None
    if not fwd:
        await update.message.reply_text(
            "⚠️ فورورد أي رسالة من قناة الكويزات هنا الأول، وبعدين ابعت /quiz_channel_id تاني."
        )
        return
    await update.message.reply_text(
        f"🆔 Quiz channel ID: <code>{fwd.id}</code>", parse_mode=ParseMode.HTML
    )

async def quiz_lectures_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """User-facing: pick a module, then a subject, then a lecture."""
    modules = ready_modules()
    if not modules:
        await update.message.reply_text("📭 مفيش محاضرات متاحة دلوقتي.")
        return
    buttons = [[InlineKeyboardButton(m, callback_data=f"module:{i}")] for i, m in enumerate(modules)]
    await update.message.reply_text(
        "📚 <b>اختار الموديول:</b>", parse_mode=ParseMode.HTML,
        reply_markup=InlineKeyboardMarkup(buttons),
    )

async def quiz_list_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Admin: numbered list of ALL lectures (open + closed) for /quiz_delete."""
    if not is_admin(update):
        await update.message.reply_text(MSG_ADMIN_ONLY)
        return
    if not QUIZ_INDEX:
        await update.message.reply_text("📭 مفيش محاضرات مسجلة لسه.")
        return
    lines = ["📋 <b>كل المحاضرات:</b>"]
    for i, (key, v) in enumerate(QUIZ_INDEX.items(), 1):
        status = "✅ مقفولة" if v["closed"] else "🟡 لسه مفتوحة"
        lecnum = f" {v['lecture_number']}" if v.get("lecture_number") else ""
        lines.append(f"{i}. {v['module']} - {v['subject']} Lecture{lecnum}: {v['name']} — {len(v['ids'])} سؤال — {status}")
    lines.append("\nاستخدم /quiz_delete &lt;رقم&gt; للحذف")
    await update.message.reply_text("\n".join(lines), parse_mode=ParseMode.HTML)

async def quiz_delete_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Admin: /quiz_delete <n> — removes a lecture from the index (does not
    delete the actual channel messages; only stops it showing up in /quiz)."""
    if not is_admin(update):
        await update.message.reply_text(MSG_ADMIN_ONLY)
        return
    if not context.args or not context.args[0].isdigit():
        await update.message.reply_text("استخدام: /quiz_delete <رقم>\nشوف الأرقام في /quiz_list")
        return
    n = int(context.args[0])
    keys = list(QUIZ_INDEX.keys())
    if n < 1 or n > len(keys):
        await update.message.reply_text(f"❌ رقم غلط — فيه {len(keys)} محاضرة بس")
        return
    key = keys[n - 1]
    removed = QUIZ_INDEX.pop(key)
    save_quiz_index()
    if QUIZ_STATE.get("current_lecture") == key:
        QUIZ_STATE["current_lecture"] = None
        save_quiz_state()
    stale_polls = [pid for pid, v in QUIZ_POLL_STATUS.items() if v["lecture"] == key]
    for pid in stale_polls:
        QUIZ_POLL_STATUS.pop(pid, None)
    save_quiz_poll_status()
    await backup_quiz_to_channel(context)
    await update.message.reply_text(
        f"🗑 اتشالت محاضرة: {removed['module']} - {removed['subject']}: {removed['name']}\n"
        "(الرسايل نفسها لسه موجودة في القناة — احذفهم يدوي لو عايز)"
    )

# ═══════════════════════════════════════════════════════════════
# TEXT MESSAGE HANDLER
# ═══════════════════════════════════════════════════════════════
async def handle(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not update.message or not update.message.text:
        return

    user_id  = update.effective_chat.id
    real_uid = update.effective_user.id if update.effective_user else user_id
    if user_id in SLEEPING:
        return

    text = update.message.text.strip()

    # ── AWAITING A QUESTION EDIT (from the review/edit prompt) ───
    pending_edit = PENDING_EDIT.pop(user_id, None)
    if pending_edit:
        items = PDF_BUFFER.get(user_id)
        idx   = pending_edit["index"]
        if not items or idx >= len(items):
            await update.message.reply_text("⚠️ السؤال ده مش موجود في البافر دلوقتي.")
            return
        item  = items[idx]
        field = pending_edit["field"]

        if field == "option":
            opt_idx = pending_edit["opt_index"]
            if 0 <= opt_idx < len(item["options"]):
                letter = string.ascii_uppercase[opt_idx]
                item["options"][opt_idx] = f"{letter}) {text}"
        elif field in ("q", "title", "content"):
            item[field] = text

        await update.message.reply_text(
            "👀 <b>راجع السؤال:</b>\n\n" + _review_text(item) + "\n\nفيه حاجة تانية عايز تعدلها؟",
            parse_mode=ParseMode.HTML,
            reply_markup=_review_buttons(idx, item),
        )
        return

    # ── AWAITING PDF NAME ────────────────────────────────────────
    if AWAITING_NAME.get(user_id):
        name = text.strip()
        PDF_NAMES[user_id]  = name
        PDF_BUFFER[user_id] = []
        PROGRESS_MSG_ID.pop(user_id, None)
        _clear_pending_image(user_id)
        _clear_clarify_queue(user_id)
        _clear_pending_edit(user_id)
        del AWAITING_NAME[user_id]
        await update.message.reply_text(
            f"📥 <b>PDF mode activated</b> — File name: <i>{name}</i>\n\n"
            "• ابعت أسئلة نصية (MCQ أو مكتوبة)\n"
            "• أو <b>فوروارد</b> كويزات أو صور/جداول مقارنة\n\n"
            "اضغط <b>Export as PDF</b> أو <b>Export as DOCX</b> لما تخلص 👇",
            parse_mode=ParseMode.HTML,
        )
        return


    # ── STORAGE PASSWORD LOOKUP ──────────────────────────────────
    if update.effective_chat.type == "private":
        items = STORAGE_INDEX.get(text.lower())
        if items:
            for message_ids in items:
                try:
                    await context.bot.copy_messages(
                        chat_id=user_id,
                        from_chat_id=STORAGE_GROUP_ID,
                        message_ids=message_ids,
                    )
                except Exception as e:
                    print(f"Storage delivery failed for password lookup: {e}")
                    await update.message.reply_text(
                        quizzy_block(QUIZZY_OOPS_ART, random.choice(QUIZZY_ERROR_LINES)),
                        parse_mode=ParseMode.HTML,
                    )
            return

    in_pdf_mode = user_id in PDF_BUFFER

    try:
        blocks = re.split(r"\n\s*\n", text)

        if not in_pdf_mode and len(blocks) > MAX_QUESTIONS_PER_MSG:
            await update.message.reply_text(
                f"❌ الحد الأقصى {MAX_QUESTIONS_PER_MSG} سؤال في المرة الواحدة"
            )
            return

        any_saved    = False
        last_label   = ""

        for block in blocks:
            block = block.strip()
            if not block:
                continue

            # ── WRITTEN ─────────────────────────────────────────
            if in_pdf_mode:
                written = parse_written_question(block)
            else:
                written = parse_written_strict(block)

            if written:
                title, content = written
                if in_pdf_mode:
                    PDF_BUFFER[user_id].append({
                        "type":    "written",
                        "title":   title,
                        "content": content,
                    })
                    any_saved  = True
                    last_label = title[:50] + ("…" if len(title) > 50 else "")
                    item_index = len(PDF_BUFFER[user_id]) - 1
                    await context.bot.send_message(
                        chat_id=update.effective_chat.id,
                        text=f"✅ اتسجل: <b>{html.escape(last_label)}</b>",
                        parse_mode=ParseMode.HTML,
                        reply_markup=_edit_button_markup(item_index),
                    )
                else:
                    await update.message.reply_text(
                        f"*{title}*\n||{content}||",
                        parse_mode=ParseMode.MARKDOWN_V2,
                    )
                continue

            # ── MCQ ─────────────────────────────────────────────
            lines = normalize_mcq_block(block)
            if len(lines) < 3:
                if not in_pdf_mode:
                    await update.message.reply_text(
                        "⚠️ <b>الصياغة غلط!</b>\n\n"
                        "الشكل الصح هو:\n"
                        "<code>السؤال\n"
                        "a) خيار 1\n"
                        "b) خيار 2 z  ← علّم الصح بـ z\n"
                        "c) خيار 3\n"
                        "ex: الشرح (اختياري)</code>",
                        parse_mode=ParseMode.HTML,
                    )
                continue

            question, raw_options, correct_index, explanation = parse_mcq_lines(lines)

            if correct_index is None or correct_index >= len(raw_options):
                if not in_pdf_mode:
                    await update.message.reply_text(
                        "⚠️ <b>ما فيش إجابة صح!</b>\n\n"
                        "علّم الإجابة الصحيحة بـ <code>z</code> في نهايتها:\n"
                        "<code>b) الإجابة الصح z</code>",
                        parse_mode=ParseMode.HTML,
                    )
                continue

            # An image sent (with no caption / unparseable caption) just
            # before this message is paired with this question, in either mode.
            pending_img = PENDING_IMAGE.pop(user_id, None)

            # ── PDF MODE ────────────────────────────────────────
            if in_pdf_mode:
                labeled_options = [
                    f"{string.ascii_uppercase[i]}) {opt}" for i, opt in enumerate(raw_options)
                ]
                item = {
                    "type": "mcq", "q": question,
                    "options": labeled_options, "correct": correct_index,
                }
                if pending_img:
                    item["image"] = pending_img
                PDF_BUFFER[user_id].append(item)
                any_saved  = True
                last_label = ("🖼 " if pending_img else "") + question[:50] + ("…" if len(question) > 50 else "")
                item_index = len(PDF_BUFFER[user_id]) - 1
                await context.bot.send_message(
                    chat_id=update.effective_chat.id,
                    text=f"✅ اتسجل: {html.escape(last_label)}",
                    parse_mode=ParseMode.HTML,
                    reply_markup=_edit_button_markup(item_index),
                )
                continue

            # ── NORMAL QUIZ MODE ─────────────────────────────────
            await deliver_quiz(
                context, user_id, question, raw_options, correct_index,
                explanation=explanation, image_path=pending_img,
            )
            events = _record_activity(real_uid, questions_delta=1)
            await react_random(update, context)
            await _announce_events(context, user_id, events)
            await backup_analytics_to_channel(context)
        if in_pdf_mode and any_saved:
            await update_progress(context, user_id, update.effective_chat.id, last_label)

    except Exception as e:
        print("ERROR:", e)

# ═══════════════════════════════════════════════════════════════
# INLINE BUTTON HANDLER
# ═══════════════════════════════════════════════════════════════
async def button_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query   = update.callback_query
    user_id = query.from_user.id
    await query.answer()

    # ── QUIZ MODULES: top-level list ──────────────────────────────
    if query.data == "quiz_modules":
        modules = ready_modules()
        if not modules:
            await query.edit_message_text("📭 مفيش محاضرات متاحة دلوقتي.")
            return
        buttons = [[InlineKeyboardButton(m, callback_data=f"module:{i}")] for i, m in enumerate(modules)]
        await query.edit_message_text(
            "📚 <b>اختار الموديول:</b>", parse_mode=ParseMode.HTML,
            reply_markup=InlineKeyboardMarkup(buttons),
        )
        return

    # ── QUIZ MODULE: list subjects within one module ───────────────
    if query.data.startswith("module:") and query.data.count(":") == 1:
        mod_idx = int(query.data.split(":")[1])
        modules = ready_modules()
        if mod_idx >= len(modules):
            await query.edit_message_text("⚠️ الموديول ده مش موجود دلوقتي.")
            return
        module = modules[mod_idx]
        subjects = ready_subjects(module)
        buttons = [
            [InlineKeyboardButton(s, callback_data=f"subject:{mod_idx}:{i}")]
            for i, s in enumerate(subjects)
        ]
        buttons.append([InlineKeyboardButton("🔙 رجوع للموديولات", callback_data="quiz_modules")])
        await query.edit_message_text(
            f"🎓 <b>{module}</b> — اختار المادة:", parse_mode=ParseMode.HTML,
            reply_markup=InlineKeyboardMarkup(buttons),
        )
        return

    # ── QUIZ SUBJECT: list lectures within one module + subject ────
    if query.data.startswith("subject:"):
        _, mod_idx_str, subj_idx_str = query.data.split(":")
        mod_idx, subj_idx = int(mod_idx_str), int(subj_idx_str)
        modules = ready_modules()
        if mod_idx >= len(modules):
            await query.edit_message_text("⚠️ الموديول ده مش موجود دلوقتي.")
            return
        module = modules[mod_idx]
        subjects = ready_subjects(module)
        if subj_idx >= len(subjects):
            await query.edit_message_text("⚠️ المادة دي مش موجودة دلوقتي.")
            return
        subject = subjects[subj_idx]
        names = ready_lecture_keys(module, subject)
        buttons = [
            [InlineKeyboardButton(
                f"Lecture {QUIZ_INDEX[name]['lecture_number'] or (i + 1)}: {QUIZ_INDEX[name]['name']}",
                callback_data=f"lecture:{mod_idx}:{subj_idx}:{i}",
            )]
            for i, name in enumerate(names)
        ]
        buttons.append([InlineKeyboardButton("🔙 رجوع للمواد", callback_data=f"module:{mod_idx}")])
        header = f"🎓 <b>{module} - {subject}</b>"
        if not names:
            header += "\n\n📭 لسه مفيش محاضرات هنا."
        await query.edit_message_text(
            header, parse_mode=ParseMode.HTML,
            reply_markup=InlineKeyboardMarkup(buttons),
        )
        return

    # ── LECTURE: start one-at-a-time delivery of a closed lecture's ready quizzes ──
    if query.data.startswith("lecture:"):
        _, mod_idx_str, subj_idx_str, lec_idx_str = query.data.split(":")
        mod_idx, subj_idx, lec_idx = int(mod_idx_str), int(subj_idx_str), int(lec_idx_str)

        modules = ready_modules()
        if mod_idx >= len(modules):
            await query.edit_message_text("⚠️ الموديول ده مش موجود دلوقتي.")
            return
        module = modules[mod_idx]
        subjects = ready_subjects(module)
        if subj_idx >= len(subjects):
            await query.edit_message_text("⚠️ المادة دي مش موجودة دلوقتي.")
            return
        subject = subjects[subj_idx]
        names = ready_lecture_keys(module, subject)
        if lec_idx >= len(names):
            await query.edit_message_text("⚠️ المحاضرة دي مش موجودة دلوقتي.")
            return
        lecture_key = names[lec_idx]
        entry = QUIZ_INDEX[lecture_key]
        ids   = entry["ids"]

        # Only polls Telegram has confirmed as stopped are actually
        # deliverable — a quiz poll's correct answer isn't known until
        # it's closed, and we need that to rebuild it as our own poll.
        closed_message_ids = {v["message_id"] for v in QUIZ_POLL_STATUS.values() if v["closed"]}
        ready_ids     = [mid for mid in ids if mid in closed_message_ids]
        not_ready_cnt = len(ids) - len(ready_ids)

        if not ready_ids:
            await query.edit_message_text(
                "⚠️ محدش قفل أي سؤال في المحاضرة دي لسه — هتتبعت لما الأدمن يوقف التصويت."
                if not_ready_cnt else
                "⚠️ المحاضرة دي مفيهاش أسئلة.",
            )
            return

        session = {
            "module": module, "subject": subject, "lecture_key": lecture_key,
            "queue": list(ready_ids), "current_poll_id": None,
            "total": len(ready_ids), "answered": 0,
        }
        LECTURE_SESSIONS[user_id] = session

        await query.edit_message_text(
            f"🎓 <b>{module} - {subject}: {entry['name']}</b> — {len(ready_ids)} سؤال، هيتبعتولك واحد واحد 👇",
            parse_mode=ParseMode.HTML,
        )

        sent = await _deliver_next_lecture_question(context, user_id, session)
        if not sent:
            LECTURE_SESSIONS.pop(user_id, None)
            await context.bot.send_message(
                chat_id=user_id,
                text="⚠️ المحاضرة دي اتحذفت من القناة، فاتشالت من القايمة.",
            )
            await backup_quiz_to_channel(context)
            return

        if not_ready_cnt:
            await context.bot.send_message(
                chat_id=user_id,
                text=(
                    f"⚠️ {not_ready_cnt} سؤال لسه مش جاهز (التصويت عليه لسه مفتوح في القناة) "
                    "— هيتبعت لما الأدمن يوقفه."
                ),
            )
        return

    # ── CLARIFY: manual correct-answer button tap ────────────────
    if query.data.startswith("clarify:"):
        _, item_index_str, choice_str = query.data.split(":")
        item_index = int(item_index_str)
        choice     = int(choice_str)

        items = PDF_BUFFER.get(user_id)
        if not items or item_index >= len(items) or items[item_index]["correct"] is not None:
            await query.edit_message_text("⚠️ السؤال ده اتحل أو اتشال بالفعل.")
            return

        item = items[item_index]
        if not (0 <= choice < len(item["options"])):
            return

        item["correct"] = choice
        POLL_WATCH.pop(item.get("poll_id"), None)

        queue = CLARIFY_QUEUE.get(user_id, [])
        if item_index in queue:
            queue.remove(item_index)

        await query.edit_message_text(
            f"✅ Q{item_index + 1}: {html.escape(item['options'][choice])}",
            parse_mode=ParseMode.HTML,
            reply_markup=_edit_button_markup(item_index),
        )

        if queue:
            await _ask_next_clarification(context, user_id, query.message.chat_id)
        else:
            CLARIFY_QUEUE.pop(user_id, None)
        return

    # ── QUESTION REVIEW / EDIT ──────────────────────────────────
    if query.data.startswith("revedit:"):
        parts      = query.data.split(":")
        item_index = int(parts[1])
        action     = parts[2]

        items = PDF_BUFFER.get(user_id)
        if not items or item_index >= len(items):
            await query.edit_message_text("⚠️ السؤال ده مش موجود في البافر دلوقتي.")
            return
        item = items[item_index]

        if action == "open":
            # First press — expand into the full edit menu
            await query.edit_message_text(
                "✏️ <b>إيه اللي عايز تعدله؟</b>\n\n" + _review_text(item),
                parse_mode=ParseMode.HTML,
                reply_markup=_review_buttons(item_index, item),
            )
            return

        if action == "done":
            await query.edit_message_text(
                "✅ <b>خلاص، اتسجل:</b>\n\n" + _review_text(item), parse_mode=ParseMode.HTML
            )
            return

        if action in ("q", "title", "content"):
            PENDING_EDIT[user_id] = {"index": item_index, "field": action}
            prompt = {
                "q":       "✏️ اكتب نص السؤال الجديد:",
                "title":   "✏️ اكتب العنوان الجديد:",
                "content": "✏️ اكتب المحتوى الجديد:",
            }[action]
            await query.edit_message_text(prompt)
            return

        if action == "opt":
            opt_idx = int(parts[3])
            if not (0 <= opt_idx < len(item["options"])):
                return
            PENDING_EDIT[user_id] = {"index": item_index, "field": "option", "opt_index": opt_idx}
            letter = string.ascii_uppercase[opt_idx]
            await query.edit_message_text(f"✏️ اكتب النص الجديد للاختيار {letter} (من غير الحرف):")
            return

        if action == "correct":
            buttons = [
                InlineKeyboardButton(string.ascii_uppercase[i], callback_data=f"revcorrect:{item_index}:{i}")
                for i in range(len(item["options"]))
            ]
            rows = [buttons[i:i + 6] for i in range(0, len(buttons), 6)]
            await query.edit_message_text("🔁 اختار الإجابة الصح:", reply_markup=InlineKeyboardMarkup(rows))
            return
        return

    if query.data.startswith("revcorrect:"):
        _, item_index_str, choice_str = query.data.split(":")
        item_index = int(item_index_str)
        choice     = int(choice_str)

        items = PDF_BUFFER.get(user_id)
        if not items or item_index >= len(items):
            await query.edit_message_text("⚠️ السؤال ده مش موجود في البافر دلوقتي.")
            return
        item = items[item_index]
        if not (0 <= choice < len(item["options"])):
            return

        item["correct"] = choice
        await query.edit_message_text(
            "👀 <b>راجع السؤال:</b>\n\n" + _review_text(item) + "\n\nفيه حاجة تانية عايز تعدلها؟",
            parse_mode=ParseMode.HTML,
            reply_markup=_review_buttons(item_index, item),
        )
        return

    # ── START MENU BUTTONS ──────────────────────────────────────
    if query.data == "menu_how":
        await query.message.reply_text(HOW_TO_USE_TEXT, parse_mode=ParseMode.HTML)
        return

    if query.data == "menu_quizzes":
        # Same as typing /quiz — sends a fresh message (not an edit) so the
        # welcome message with its buttons stays intact above it.
        modules = ready_modules()
        if not modules:
            await query.message.reply_text("📭 مفيش محاضرات متاحة دلوقتي.")
            return
        buttons = [[InlineKeyboardButton(m, callback_data=f"module:{i}")] for i, m in enumerate(modules)]
        await query.message.reply_text(
            "📚 <b>اختار الموديول:</b>", parse_mode=ParseMode.HTML,
            reply_markup=InlineKeyboardMarkup(buttons),
        )
        return

    # ── EXPORT BUTTONS ──────────────────────────────────────────
    items = PDF_BUFFER.get(user_id, [])
    name  = PDF_NAMES.get(user_id, "questions")
    safe  = re.sub(r"[^\w\s\-]", "", name).strip().replace(" ", "_") or "questions"

    if query.data == "gen_pdf":
        if not items:
            await query.message.reply_text(MSG_EXPORT_EMPTY)
            return
        await query.message.reply_text(MSG_EXPORT_GENERATING.format(kind="PDF", count=len(items)))
        await _export_pdf_session(context, query.message, user_id, user_id, items, name, fmt="pdf")

    elif query.data == "gen_docx":
        if not items:
            await query.message.reply_text(MSG_EXPORT_EMPTY)
            return
        await query.message.reply_text(MSG_EXPORT_GENERATING.format(kind="DOCX", count=len(items)))
        await _export_pdf_session(context, query.message, user_id, user_id, items, name, fmt="docx")

    elif query.data == "clear_pdf":
        _reset_pdf_session(user_id)
        await query.message.reply_text(MSG_EXPORT_CLEARED_ALL)

# ═══════════════════════════════════════════════════════════════
# PDF/DOCX EXPORT — single source of truth, called from both the
# /pdf_generate and /pdf_clear commands and their button equivalents
# (gen_pdf/gen_docx/clear_pdf), so there's only one place that can
# forget to track XP or diverge in behavior between the two.
# ═══════════════════════════════════════════════════════════════
def _reset_pdf_session(user_id: int) -> None:
    """Clears everything tied to an in-progress PDF-collection session —
    used after a successful export and by explicit clear/cancel alike."""
    _cleanup_images(user_id)
    _clear_pending_image(user_id)
    _clear_clarify_queue(user_id)
    _clear_pending_edit(user_id)
    PDF_BUFFER.pop(user_id, None)
    PDF_NAMES.pop(user_id, None)
    AWAITING_NAME.pop(user_id, None)
    PROGRESS_MSG_ID.pop(user_id, None)

async def _export_pdf_session(context: ContextTypes.DEFAULT_TYPE, message, session_id: int,
                               analytics_uid: int, items: list, name: str, fmt: str) -> bool:
    """Builds a PDF or DOCX (fmt='pdf'|'docx') from items, sends it via
    message.reply_document, records the export XP/counter against
    analytics_uid, announces any level-up, and resets the session keyed by
    session_id. Returns False (having already replied with the reason) if
    DOCX isn't available or the build blew up."""
    safe = re.sub(r"[^\w\s\-]", "", name).strip().replace(" ", "_") or "questions"

    if fmt == "docx":
        if not DOCX_AVAILABLE:
            await message.reply_text(MSG_DOCX_UNAVAILABLE)
            return False
        try:
            doc_bytes = build_docx(items, name)
        except Exception as e:
            print("DOCX ERROR:", e)
            await message.reply_text(
                f"{quizzy_block(QUIZZY_OOPS_ART, random.choice(QUIZZY_ERROR_LINES))}\n\n<code>{e}</code>",
                parse_mode=ParseMode.HTML,
            )
            return False
        await message.reply_document(
            document=doc_bytes, filename=f"{safe}.docx",
            caption=MSG_DOCX_CAPTION.format(count=len(items), name=name, quizzy_line=random.choice(QUIZZY_SUCCESS_LINES)),
            parse_mode=ParseMode.HTML,
        )
    else:
        pdf_bytes = build_pdf(items, name)
        await message.reply_document(
            document=pdf_bytes, filename=f"{safe}.pdf",
            caption=MSG_PDF_CAPTION.format(count=len(items), name=name, quizzy_line=random.choice(QUIZZY_SUCCESS_LINES)),
            parse_mode=ParseMode.HTML,
        )

    q_count = sum(1 for it in items if it.get("type") in ("mcq", "written"))
    events  = _record_activity(analytics_uid, questions_delta=q_count, pdfs_delta=1, session_questions=q_count)
    await _announce_events(context, session_id, events)
    await backup_analytics_to_channel(context)
    _reset_pdf_session(session_id)
    return True

# ═══════════════════════════════════════════════════════════════
# PDF COMMANDS
# ═══════════════════════════════════════════════════════════════
async def pdf_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_chat.id
    _reset_pdf_session(user_id)
    AWAITING_NAME[user_id] = True
    await update.message.reply_text(MSG_PDF_ASK_NAME, parse_mode=ParseMode.HTML)

async def pdf_generate(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id  = update.effective_chat.id
    real_uid = update.effective_user.id if update.effective_user else user_id
    items    = PDF_BUFFER.get(user_id, [])
    if not items:
        await update.message.reply_text(MSG_PDF_EMPTY)
        return
    name = PDF_NAMES.get(user_id, "questions")
    await update.message.reply_text(MSG_PDF_GENERATING.format(count=len(items)))
    await _export_pdf_session(context, update.message, user_id, real_uid, items, name, fmt="pdf")

async def pdf_clear(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_chat.id
    _reset_pdf_session(user_id)
    await update.message.reply_text(MSG_PDF_CLEARED)

async def cancel_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Bails out of whatever's in progress: PDF collection session or a
    pending image waiting for its question."""
    user_id = update.effective_chat.id
    was_doing_something = bool(
        PDF_BUFFER.get(user_id) or AWAITING_NAME.get(user_id)
        or PENDING_IMAGE.get(user_id)
    )
    _cleanup_images(user_id)
    _clear_pending_image(user_id)
    _clear_clarify_queue(user_id)
    _clear_pending_edit(user_id)
    PDF_BUFFER.pop(user_id, None)
    PDF_NAMES.pop(user_id, None)
    AWAITING_NAME.pop(user_id, None)
    PROGRESS_MSG_ID.pop(user_id, None)
    if was_doing_something:
        await update.message.reply_text(MSG_CANCEL_DONE)
    else:
        await update.message.reply_text(MSG_CANCEL_NOTHING)

# ═══════════════════════════════════════════════════════════════
# START  (also wakes bot from sleep)
# ═══════════════════════════════════════════════════════════════
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    chat_id = update.effective_chat.id
    SLEEPING.discard(chat_id)

    if chat_id not in USERS:
        USERS.add(chat_id)
        save_users()
        await backup_storage_to_channel(context)

    await update.message.reply_text(
        f"{quizzy_block(QUIZZY_WELCOME_ART, random.choice(QUIZZY_WELCOME_LINES))}\n\n"
        "تحب تعمل أي؟!:",
        parse_mode=ParseMode.HTML,
        reply_markup=start_menu_keyboard(),
    )

async def commands_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """/c — lists every command, admin-only ones only shown to the admin."""
    lines = ["📖 <b>الأوامر المتاحة:</b>\n"]
    lines.append("👤 <b>للجميع</b>")
    lines.append("/start — القائمة الرئيسية")
    lines.append("/sleep — يوقف البوت مؤقتًا في المحادثة دي")
    lines.append("/mystats — إحصائياتك (أسئلة أنشأتها، سلسلة الأيام)")
    lines.append("/pdf_start — يبدأ سيشن تجميع صور لملف PDF")
    lines.append("/pdf_generate — يطلع PDF من الصور اللي جمعتها")
    lines.append("/pdf_clear — يمسح سيشن الـ PDF الحالي")
    lines.append("/cancel — يلغي أي حاجة شغالة دلوقتي (PDF، صورة معلّقة، إلخ)")
    lines.append("/quiz — تصفح المحاضرات (موديول ← مادة ← محاضرة) وسحب أسئلتها")
    lines.append("/storage_id — يجيب chat ID بتاع المكان ده (لضبط STORAGE_GROUP_ID)")
    lines.append("/quiz_channel_id — يجيب chat ID لقناة الكويز (فوروارد رسالة منها الأول)")
    lines.append("🤖 ابعت صورة/PDF أسئلة من غير كابشن — هيعرض زرار استخراج بالـ AI")
    lines.append("/c — القائمة دي")

    if is_admin(update):
        lines.append("\n🔐 <b>للأدمن بس</b>")
        lines.append("/admincheck — يتأكد إنك أدمن")
        lines.append("/broadcast &lt;رسالة&gt; — يبعت رسالة لكل المستخدمين")
        lines.append("/backup_now — يعمل ريفريش فوري للباك أب المثبّت (ستوريدج + كويز)")
        lines.append("/quiz_list — ليستة مرقّمة بكل المحاضرات (مفتوحة ومقفولة)")
        lines.append("/quiz_delete &lt;رقم&gt; — يشيل محاضرة من الفهرس")

    await update.message.reply_text("\n".join(lines), parse_mode=ParseMode.HTML)

# ═══════════════════════════════════════════════════════════════
# ADMIN HELPERS
# ═══════════════════════════════════════════════════════════════
def is_admin(update: Update) -> bool:
    return update.effective_user and update.effective_user.id == ADMIN_ID

async def admincheck_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    uid = update.effective_user.id if update.effective_user else "?"
    if is_admin(update):
        await update.message.reply_text(
            f"✅ <b>أنت الأدمن!</b>\n"
            f"🆔 Your ID: <code>{uid}</code>\n"
            f"👥 Total users: <b>{len(USERS)}</b>",
            parse_mode=ParseMode.HTML,
        )
    else:
        await update.message.reply_text(
            f"🚫 مش أدمن\n🆔 Your ID: <code>{uid}</code>",
            parse_mode=ParseMode.HTML,
        )

# ═══════════════════════════════════════════════════════════════
# BROADCAST COMMAND  (admin only)
# ═══════════════════════════════════════════════════════════════
async def broadcast_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update):
        await update.message.reply_text(MSG_ADMIN_ONLY)
        return

    # Message text comes after /broadcast, or from a replied-to message
    if context.args:
        text = " ".join(context.args)
    elif update.message.reply_to_message and update.message.reply_to_message.text:
        text = update.message.reply_to_message.text
    else:
        await update.message.reply_text(
            "⚠️ استخدام:\n"
            "<code>/broadcast رسالتك هنا</code>\n\n"
            "أو رد بـ /broadcast على رسالة موجودة.",
            parse_mode=ParseMode.HTML,
        )
        return

    if not text.strip():
        await update.message.reply_text("❌ الرسالة فارغة")
        return

    users_list = list(USERS)
    total      = len(users_list)

    status_msg = await update.message.reply_text(
        f"📡 <b>جاري الإرسال لـ {total} مستخدم...</b>",
        parse_mode=ParseMode.HTML,
    )

    success = 0
    failed  = 0
    blocked = []

    for uid in users_list:
        try:
            await context.bot.send_message(
                chat_id=uid,
                text=text,
                parse_mode=ParseMode.HTML,
            )
            success += 1
        except Forbidden:
            # The user actually blocked the bot (or deleted their account) —
            # this is the only case where removing them from USERS is safe.
            failed += 1
            blocked.append(uid)
        except Exception as e:
            # Any other error (network blip, rate limit, Telegram hiccup) is
            # NOT proof the user blocked us — keep them in USERS so a
            # transient failure doesn't silently and permanently unsubscribe
            # a real, still-active user.
            failed += 1
            print(f"Broadcast failed for {uid} (not removed — not a block):", e)

    # Remove users who blocked the bot
    if blocked:
        for uid in blocked:
            USERS.discard(uid)
        save_users()
        await backup_storage_to_channel(context)

    summary = (
        f"✅ <b>Broadcast اتبعت!</b>\n\n"
        f"👥 المستخدمين: <b>{total}</b>\n"
        f"✔️ نجح: <b>{success}</b>\n"
        f"❌ فشل / بلوك: <b>{failed}</b>"
    )
    if blocked:
        summary += f"\n🗑 تم حذف {len(blocked)} يوزر بلوك البوت من القائمة"

    await status_msg.edit_text(summary, parse_mode=ParseMode.HTML)

# ═══════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════
async def mystats_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    # Analytics are keyed by the person's real Telegram user id — not the
    # chat id — so this shows the same numbers whether /mystats is run in
    # a DM or inside a group/channel the bot is in.
    user_id = update.effective_user.id if update.effective_user else update.effective_chat.id
    entry   = ANALYTICS.get(str(user_id))
    if not entry or not entry.get("last_active_date"):
        await update.message.reply_text("📊 لسه معندكش إحصائيات. ابعت أسئلة وهتظهر هنا!")
        return

    streak  = entry.get("streak", 0)
    total_q = entry.get("questions_created", 0)
    pdfs    = entry.get("pdfs_exported", 0)
    last    = entry.get("last_active_date", "—")
    xp      = entry.get("xp", 0)
    level   = entry.get("level", 0)
    title   = _level_title(level)
    flame   = "🔥" * min(streak, 5) if streak else "❄️"

    xp_start, xp_end = _level_xp_range(level)
    xp_into_level    = xp - xp_start
    xp_needed        = xp_end - xp_start
    bar_filled       = int((xp_into_level / xp_needed) * 10) if xp_needed else 10
    bar              = "█" * bar_filled + "░" * (10 - bar_filled)

    # achievements summary
    ach        = entry.get("achievements", {})
    ach_lines  = []
    icons      = {"questions": "❓", "streak": "🔥", "pdfs": "📚", "speed": "⚡"}
    tier_names = ["", "I", "II", "III", "IV", "V"]
    for key, emoji in icons.items():
        tier = ach.get(key, 0)
        if tier:
            name = ACHIEVEMENTS[key][tier - 1][1]
            ach_lines.append(f"  {emoji} {name} {'⭐' * tier}")

    ach_text = "\n".join(ach_lines) if ach_lines else "  لسه مفيش إنجازات"

    await update.message.reply_text(
        f"📊 <b>إحصائياتك</b>\n\n"
        f"🏅 المستوى: <b>{level}</b> — <i>{title}</i>\n"
        f"✨ XP: <b>{xp}</b>  [{bar}]  → {xp_end}\n\n"
        f"❓ أسئلة أنشأتها: <b>{total_q}</b>\n"
        f"📚 PDFs: <b>{pdfs}</b>\n"
        f"🗓 سلسلة الأيام: <b>{streak}</b> {flame}\n"
        f"📅 آخر نشاط: <b>{last}</b>\n\n"
        f"🏆 <b>إنجازات:</b>\n{ach_text}",
        parse_mode=ParseMode.HTML,
    )


async def reset_analytics_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Admin-only: wipe all analytics data locally and delete the pinned
    backup in the analytics group. Use when you need a clean slate."""
    global _analytics_backup_msg_id
    if update.effective_chat.id != ADMIN_ID:
        return
    ANALYTICS.clear()
    save_analytics()
    if _analytics_backup_msg_id and ANALYTICS_GROUP_ID:
        try:
            await context.bot.delete_message(
                chat_id=ANALYTICS_GROUP_ID,
                message_id=_analytics_backup_msg_id,
            )
        except Exception:
            pass
    _analytics_backup_msg_id = None
    await update.message.reply_text("🗑 Analytics wiped — local file cleared and backup deleted.")

async def restore_analytics_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Admin-only: manually re-pull analytics.json from the pinned backup
    in ANALYTICS_GROUP_ID — the same restore that runs automatically on
    startup. Use this if the local file ever gets wiped, corrupted, or
    out of sync with the backup, without needing to restart the bot."""
    if not (update.effective_user and update.effective_user.id == ADMIN_ID):
        return
    before = len(ANALYTICS)
    await restore_analytics_from_channel(context)
    await update.message.reply_text(
        f"♻️ Restored from pinned backup.\n"
        f"Users on file: <b>{len(ANALYTICS)}</b> (was {before} before restore).",
        parse_mode=ParseMode.HTML,
    )

async def import_analytics_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Admin-only: import analytics from an uploaded .json file. Reply to
    the file's message with /import_analytics. This is the fallback for
    when the pinned backup itself is missing/corrupted — e.g. importing a
    copy you saved elsewhere. Merges into (does not wipe) existing data,
    then re-saves and re-pins so the channel backup reflects the import."""
    if not (update.effective_user and update.effective_user.id == ADMIN_ID):
        return
    reply = update.message.reply_to_message
    doc   = reply.document if reply else None
    if not doc:
        await update.message.reply_text(
            "⚠️ Reply to the analytics .json file with /import_analytics."
        )
        return
    try:
        tg_file  = await context.bot.get_file(doc.file_id)
        raw      = await tg_file.download_as_bytearray()
        imported = json.loads(bytes(raw).decode("utf-8"))
        if not isinstance(imported, dict):
            raise ValueError("File doesn't look like an analytics export (expected a JSON object).")
    except Exception as e:
        await update.message.reply_text(f"❌ Import failed: {e}")
        return
    ANALYTICS.update(imported)
    save_analytics()
    await backup_analytics_to_channel(context)
    await update.message.reply_text(
        f"✅ Imported <b>{len(imported)}</b> user(s) — merged into current data, "
        f"saved locally, and re-pinned to the backup channel.\n"
        f"Total users on file now: <b>{len(ANALYTICS)}</b>.",
        parse_mode=ParseMode.HTML,
    )

async def _post_init(app):
    """Runs once after the bot connects, before polling starts — restores
    the storage-group and quiz-channel indexes from their pinned backup
    messages, so a wiped/switched local disk doesn't orphan content that's
    still sitting safely in the channels themselves."""
    await restore_storage_from_channel(app)
    await restore_quiz_from_channel(app)
    await restore_analytics_from_channel(app)

app = ApplicationBuilder().token(BOT_TOKEN).post_init(_post_init).build()

app.add_handler(CommandHandler("start",          start))
app.add_handler(CommandHandler("c",              commands_cmd))
app.add_handler(CommandHandler("cancel",         cancel_cmd))
app.add_handler(CommandHandler("sleep",          sleep_cmd))
app.add_handler(CommandHandler("admincheck",     admincheck_cmd))
app.add_handler(CommandHandler("broadcast",      broadcast_cmd))
app.add_handler(CommandHandler("pdf_start",      pdf_start))
app.add_handler(CommandHandler("pdf_generate",   pdf_generate))
app.add_handler(CommandHandler("pdf_clear",      pdf_clear))
# Storage group setup helper
app.add_handler(CommandHandler("mystats",           mystats_cmd))
app.add_handler(CommandHandler("restore_analytics", restore_analytics_cmd))
app.add_handler(CommandHandler("import_analytics",  import_analytics_cmd))
app.add_handler(CommandHandler("reset_analytics",   reset_analytics_cmd))
app.add_handler(CommandHandler("storage_id",     storage_id_cmd))
app.add_handler(CommandHandler("backup_now",     backup_now_cmd))
# Quiz channel
app.add_handler(CommandHandler("quiz_channel_id", quiz_channel_id_cmd))
app.add_handler(CommandHandler("quiz",            quiz_lectures_cmd))
app.add_handler(CommandHandler("quiz_list",       quiz_list_cmd))
app.add_handler(CommandHandler("quiz_delete",     quiz_delete_cmd))

# Poll handler before text handler (forwarded OR own quiz polls) —
# excludes the quiz channel, which has its own dedicated handler below.
app.add_handler(MessageHandler(filters.POLL & ~filters.Chat(QUIZ_CHANNEL_ID), handle_poll))

# Quiz channel indexing — lecture titles, "-END", and quiz polls posted
# there get filed by handle_quiz_channel_message, not treated as a user's
# own quiz-building activity. Must be registered before the generic
# text/poll handlers below.
app.add_handler(MessageHandler(
    filters.Chat(QUIZ_CHANNEL_ID) & (filters.POLL | filters.TEXT), handle_quiz_channel_message
))

# Storage group indexing — anything posted in the vault group gets filed by
# its caption's password word. Must be checked before the generic photo
# handler below so vault posts don't get mistaken for quiz images.
STORAGE_MEDIA_FILTER = (
    filters.PHOTO | filters.VIDEO | filters.Document.ALL
    | filters.AUDIO | filters.VOICE | filters.ANIMATION | filters.Sticker.ALL
)
app.add_handler(MessageHandler(
    filters.Chat(STORAGE_GROUP_ID) & STORAGE_MEDIA_FILTER, handle_storage_message
))

# Image handler (photos in PDF mode) — excludes the storage group
app.add_handler(MessageHandler(
    filters.PHOTO & ~filters.Chat(STORAGE_GROUP_ID), handle_image
))

# PDF handler — captioned PDFs in a private DM are parsed as manual MCQs;
# excludes the storage group and quiz channel.
app.add_handler(MessageHandler(
    filters.Document.PDF & ~filters.Chat(STORAGE_GROUP_ID) & ~filters.Chat(QUIZ_CHANNEL_ID), handle_document
))

# Inline buttons
app.add_handler(CallbackQueryHandler(button_handler))
app.add_handler(PollHandler(poll_update_handler))
app.add_handler(PollAnswerHandler(handle_poll_answer))

# Text handler last — excludes the storage group and the quiz channel
app.add_handler(MessageHandler(
    filters.TEXT & ~filters.COMMAND & ~filters.Chat(STORAGE_GROUP_ID) & ~filters.Chat(QUIZ_CHANNEL_ID), handle
))

print("Bot running... V5.5")
app.run_polling()
